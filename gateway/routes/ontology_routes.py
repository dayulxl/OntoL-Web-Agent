"""
本体建模 API 路由
----------------
提供图数据库（Memgraph/Neo4j）的节点/关系 CRUD、Schema 发现和图遍历接口。
"""
from typing import Optional
from pathlib import Path as _Path
import json

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from pydantic import BaseModel, Field

from common.exceptions.base import InfrastructureError

router = APIRouter(tags=["Ontology"])


# =========================================================================
# Pydantic 模型
# =========================================================================

class NodeCreate(BaseModel):
    label: str = Field(..., description="节点标签", examples=["Entity"])
    properties: dict = Field(default_factory=dict, description="节点属性", examples=[{"name": "航母战斗群", "type": "海军编队"}])

class NodeUpdate(BaseModel):
    properties: dict = Field(..., description="要更新的属性")

class EdgeCreate(BaseModel):
    source_id: int = Field(..., description="起始节点 ID")
    target_id: int = Field(..., description="目标节点 ID")
    rel_type: str = Field(..., description="关系类型", examples=["DEPLOYED_TO"])
    properties: Optional[dict] = Field(default_factory=dict, description="关系属性")

class CypherQuery(BaseModel):
    query: str = Field(..., description="只读 Cypher 语句")
    params: Optional[dict] = Field(default_factory=dict, description="查询参数")


class ToolsCallBody(BaseModel):
    name: str = Field(..., description="工具名称（如 infer_forward）")
    arguments: dict = Field(default_factory=dict, description="工具参数")


# ── ontol_model / ontol_model_attr 请求体 ──

class OntolModelCreateBody(BaseModel):
    model_config = {"extra": "ignore"}

    code: str = Field(..., max_length=50, description="本体编码（唯一）")
    ontol_parent_id: Optional[str] = Field(None, max_length=32, description="父级模型ID")
    name: str = Field(..., max_length=50, description="本体名称")
    ontol_model_type: str = Field(..., max_length=2, description="本体类型：M1/M2/M3/M4/M5/M6/M7/ME/MT")
    ontol_model_status: str = Field("0", max_length=2, description="本体状态：0=启用中 1=已停用")
    ontol_model_desc: Optional[str] = Field(None, max_length=255, description="本体描述")

class OntolModelUpdateBody(BaseModel):
    model_config = {"extra": "ignore"}
    code: Optional[str] = Field(None, max_length=50)
    name: Optional[str] = Field(None, max_length=50)
    ontol_model_type: Optional[str] = Field(None, max_length=2)
    ontol_model_status: Optional[str] = Field(None, max_length=2)
    ontol_model_desc: Optional[str] = Field(None, max_length=255)

class OntolModelAttrCreateBody(BaseModel):
    model_config = {"extra": "ignore"}

    id: str = Field(..., max_length=32, description="属性ID（主键）")
    ontol_model_id: Optional[str] = Field(None, max_length=32)
    name: str = Field(..., max_length=50)
    code: str = Field(..., max_length=50)
    attr_data_type: str = Field("VARCHAR", max_length=20)
    attr_length: Optional[str] = Field(None, max_length=10)
    attr_digit: Optional[str] = Field(None, max_length=10)
    attr_is_only: Optional[str] = Field("0", max_length=2)
    attr_cascade_colum: Optional[str] = Field(None, max_length=50)
    attr_data_source_flag: Optional[str] = Field(None, max_length=2)
    attr_data_source: Optional[str] = Field(None, max_length=255)
    attr_required: Optional[str] = Field("0", max_length=2)
    attr_default_value: Optional[str] = Field(None, max_length=500)
    attr_is_system: Optional[str] = Field("0", max_length=2)
    attr_desc: Optional[str] = Field(None, max_length=50)

class OntolModelAttrUpdateBody(BaseModel):
    model_config = {"extra": "ignore"}

    name: Optional[str] = Field(None, max_length=50)
    code: Optional[str] = Field(None, max_length=50)
    attr_data_type: Optional[str] = Field(None, max_length=20)
    attr_length: Optional[str] = Field(None, max_length=10)
    attr_digit: Optional[str] = Field(None, max_length=10)
    attr_is_only: Optional[str] = Field(None, max_length=2)
    attr_cascade_colum: Optional[str] = Field(None, max_length=50)
    attr_data_source_flag: Optional[str] = Field(None, max_length=2)
    attr_data_source: Optional[str] = Field(None, max_length=255)
    attr_required: Optional[str] = Field(None, max_length=2)
    attr_default_value: Optional[str] = Field(None, max_length=500)
    attr_is_system: Optional[str] = Field(None, max_length=2)
    attr_desc: Optional[str] = Field(None, max_length=50)


# =========================================================================
# 依赖注入
# =========================================================================

async def get_graph():
    """
    获取 GraphMemory 实例（惰性导入，避免启动时图数据库未就绪而崩溃）。
    """
    from infrastructure.db.neo4j import get_driver
    from capabilities.memory.graph_memory import GraphMemory

    try:
        driver = await get_driver()
    except InfrastructureError:
        raise HTTPException(status_code=503, detail="Graph DB driver not initialized")
    return GraphMemory(driver)

async def get_ontology_repo():
    """获取 OntologyRepo 实例。"""
    from infrastructure.db.sqlite_db import get_sqlite_pool
    from infrastructure.db.ontology_repo import OntologyRepo
    try:
        pool = get_sqlite_pool()
        return OntologyRepo(pool)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Database not available: {e}")


# =========================================================================
# 历史记录辅助
# =========================================================================

async def _record_history(node_id: str, action: str, context: dict, create_id: str = ""):
    """将图数据变更写入 ontol_data_his 历史表。"""
    import uuid as _uuid
    record = {
        "id": _uuid.uuid4().hex[:16],
        "node_id": str(node_id),
        "action": action,
        "timestamp": datetime.utcnow().isoformat() + "Z",
        **context,
    }
    try:
        await _execute_scene(
            "INSERT INTO ontol_data_his (id, node_id, context, create_id) VALUES (?,?,?,?)",
            (record["id"], str(node_id), json.dumps(record, ensure_ascii=False), create_id),
        )
    except Exception:
        pass  # 历史记录失败不影响主流程


async def _bump_node_version(graph, node_id: int, node_name: str = ""):
    """递增图节点版本号。"""
    try:
        from infrastructure.db.neo4j import get_driver
        driver = await get_driver()
        async with driver.session() as session:
            # 读取当前版本号
            result = await session.run(
                "MATCH (n) WHERE id(n) = $node_id "
                "RETURN coalesce(n.version, '0') AS ver, n.name AS name",
                node_id=node_id,
            )
            rec = await result.single()
            if rec:
                old_ver = int(rec["ver"]) if rec["ver"] and str(rec["ver"]).isdigit() else 0
                new_ver = str(old_ver + 1)
                await session.run(
                    "MATCH (n) WHERE id(n) = $node_id SET n.version = $ver",
                    node_id=node_id, ver=new_ver,
                )
    except Exception:
        pass  # 版本号失败不影响主流程


# =========================================================================
# Schema
# =========================================================================

@router.get("/ontology/schema")
async def ontology_schema(graph=Depends(get_graph)):
    """获取图 Schema：所有标签、关系类型、节点和边计数。"""
    try:
        return await graph.get_schema()
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB query failed: {e}")


# =========================================================================
# Neighborhood — 图邻域查询
# =========================================================================

@router.get("/ontology/neighborhood/{node_id}")
async def get_neighborhood(node_id: int, depth: int = 1, graph=Depends(get_graph)):
    """获取节点的图邻域（关联节点 + 关系），支持 depth 1-3。"""
    try:
        return await graph.get_neighborhood(node_id, depth=depth)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB query failed: {e}")


# =========================================================================
# Node CRUD
# =========================================================================

@router.get("/ontology/nodes")
async def list_nodes(
    label: Optional[str] = None,
    limit: int = 100,
    keyword: Optional[str] = None,
    graph=Depends(get_graph),
):
    """列出节点，支持按标签和关键词过滤。"""
    try:
        return await graph.list_nodes(label=label, limit=min(limit, 500), keyword=keyword)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB query failed: {e}")


@router.get("/ontology/nodes/{node_id}/history")
async def get_node_history(node_id: int):
    """获取图节点的历史变更记录（ontol_data_his）。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_data_his WHERE node_id=? AND delete_flag='0' ORDER BY create_time DESC LIMIT 50",
            (str(node_id),),
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"History query failed: {e}")


@router.get("/ontology/nodes/{node_id}")
async def get_node(node_id: int, graph=Depends(get_graph)):
    """获取节点详情及邻接关系。"""
    try:
        node = await graph.get_node(node_id)
        if node is None:
            raise HTTPException(status_code=404, detail=f"Node {node_id} not found")
        return node
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB query failed: {e}")


@router.post("/ontology/nodes", status_code=201)
async def create_node(body: NodeCreate, graph=Depends(get_graph)):
    """创建节点，自动生成雪花 ID（64 位纯数字）。"""
    try:
        # 生成雪花 ID（去重）
        import time as _time
        sf = SnowflakeGenerator(
            worker_id=(int(_time.time() * 1000) & 0x1F),
            datacenter_id=1,
        )
        # 查询已存在的节点 ID，避免冲突
        existing = await graph.execute_readonly_cypher("MATCH (n) WHERE n.id IS NOT NULL RETURN n.id AS id LIMIT 5000")
        existing_ids = {r["id"] for r in existing if isinstance(r.get("id"), int)}
        snowflake_id = sf.next_id()
        while snowflake_id in existing_ids:
            snowflake_id = sf.next_id()

        props = dict(body.properties)
        # 雪花 ID：如果前端没传 id 或 id 为空/占位符，自动生成
        if not props.get("id") or props.get("id") == "大模型随机生成":
            props["id"] = snowflake_id

        result = await graph.create_node(body.label, props)
        # 记录历史
        await _record_history(str(result.get("id", "")), "create", {"label": body.label, "new_props": props})
        return result
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB create failed: {e}")


@router.put("/ontology/nodes/{node_id}")
async def update_node(node_id: int, body: NodeUpdate, graph=Depends(get_graph)):
    """更新节点属性。自动记录历史并递增版本号。删除的属性会被移除。"""
    try:
        # 先获取旧状态
        old_node = await graph.get_node(node_id)
        old_props = old_node["properties"] if old_node else {}
        new_props = body.properties

        # 找出被删除的 key（旧有但新无）
        removed_keys = [k for k in old_props if k not in new_props]
        node = await graph.update_node(node_id, new_props, remove_keys=removed_keys or None)
        if node is None:
            raise HTTPException(status_code=404, detail=f"Node {node_id} not found")

        # 记录历史（变更前后的差异）
        await _record_history(str(node_id), "update", {
            "old_props": old_props,
            "new_props": new_props,
            "removed_keys": removed_keys,
        })
        # 递增版本号
        await _bump_node_version(graph, node_id)
        return node
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB update failed: {e}")


@router.delete("/ontology/nodes/{node_id}")
async def delete_node(node_id: int, graph=Depends(get_graph)):
    """删除节点及其所有关系。删除前记录历史快照。"""
    try:
        # 删除前获取快照
        snapshot = await graph.get_node(node_id)
        if snapshot:
            await _record_history(str(node_id), "delete", {
                "snapshot": {"labels": snapshot["labels"], "properties": snapshot["properties"]},
            })
        ok = await graph.delete_node(node_id)
        if not ok:
            raise HTTPException(status_code=404, detail=f"Node {node_id} not found")
        return {"deleted": True, "node_id": node_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB delete failed: {e}")


# =========================================================================
# Edge CRUD
# =========================================================================

@router.post("/ontology/edges", status_code=201)
async def create_edge(body: EdgeCreate, graph=Depends(get_graph)):
    """创建关系。记录历史并递增源/目标节点版本号。"""
    try:
        result = await graph.create_edge(body.source_id, body.target_id, body.rel_type, body.properties)
        # 记录历史
        await _record_history(str(result.get("source_id", "")), "edge_create", {
            "target_id": body.target_id,
            "rel_type": body.rel_type,
            "edge_props": body.properties or {},
            "edge_id": result.get("id", ""),
        })
        # 递增双方版本号
        await _bump_node_version(graph, body.source_id)
        await _bump_node_version(graph, body.target_id)
        return result
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB create edge failed: {e}")


@router.delete("/ontology/edges/{edge_id}")
async def delete_edge(edge_id: int, graph=Depends(get_graph)):
    """删除关系。记录历史快照并递增版本号。"""
    try:
        # 删除前获取边信息用于记录
        from infrastructure.db.neo4j import get_driver
        source_id = 0; target_id = 0
        try:
            driver = await get_driver()
            async with driver.session() as session:
                r = await session.run(
                    "MATCH ()-[e]->() WHERE id(e) = $eid "
                    "RETURN type(e) AS rel_type, id(startNode(e)) AS src, id(endNode(e)) AS tgt",
                    eid=edge_id,
                )
                rec = await r.single()
                if rec:
                    source_id = rec["src"]
                    target_id = rec["tgt"]
                    await _record_history(str(source_id), "edge_delete", {
                        "target_id": target_id,
                        "rel_type": rec["rel_type"],
                        "edge_id": edge_id,
                    })
        except Exception:
            pass  # 记录失败不影响删除

        ok = await graph.delete_edge(edge_id)
        if not ok:
            raise HTTPException(status_code=404, detail=f"Edge {edge_id} not found")

        # 递增版本号
        if source_id: await _bump_node_version(graph, source_id)
        if target_id: await _bump_node_version(graph, target_id)
        return {"deleted": True, "edge_id": edge_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB delete edge failed: {e}")


class EdgeUpdate(BaseModel):
    properties: dict = Field(default_factory=dict, description="边的属性（完整替换）")


@router.put("/ontology/edges/{edge_id}")
async def update_edge(edge_id: int, body: EdgeUpdate, graph=Depends(get_graph)):
    """更新边的属性。"""
    try:
        result = await graph.update_edge(edge_id, body.properties or {})
        if not result:
            raise HTTPException(status_code=404, detail=f"Edge {edge_id} not found")
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB update edge failed: {e}")


@router.get("/ontology/edges")
async def list_edges(limit: int = 2000, graph=Depends(get_graph)):
    """列出所有关系（边）。"""
    try:
        return await graph.list_all_edges(limit=min(limit, 5000))
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB query failed: {e}")


@router.post("/tools/call")
async def tools_call(body: ToolsCallBody):
    """代理转发到 KG 推理机的 /tools/call。"""
    import json, requests as req
    from common.config.settings import get_settings

    kg_url = get_settings().kg_server_url
    try:
        resp = req.post(
            f"{kg_url}/tools/call",
            json={"name": body.name, "arguments": body.arguments},
            timeout=120,
        )
        # 把 KG 服务端的响应原样返回，状态码也透传
        try:
            data = resp.json()
        except json.JSONDecodeError:
            data = {"raw": resp.text}
        from fastapi.responses import JSONResponse
        return JSONResponse(content=data, status_code=resp.status_code)
    except req.RequestException as e:
        raise HTTPException(status_code=503, detail=f"KG reasoning server unreachable: {e}")


class InferOnNodesBody(BaseModel):
    node_ids: list[str] = Field(default_factory=list, description="节点内置ID列表")
    confidence: float = Field(default=0.8, description="置信度")
    copy_version: str = Field(default="", alias="cope_version", description="副本版本号")

    class Config:
        populate_by_name = True


@router.post("/infer-on-nodes")
async def infer_on_nodes(body: InferOnNodesBody):
    """调用内部图推理机 — 流式收集结果，返回 messages 数组给前端。"""
    from business.reasoning import run_reasoning

    # 解析 node_id：支持纯数字 Snowflake ID 或字符串 code
    async def _resolve_seed_id(raw: str) -> int | None:
        try:
            return int(raw)
        except ValueError:
            pass
        # 可能是 code 或名称
        try:
            from infrastructure.db.neo4j import get_driver
            driver = await get_driver()
            async with driver.session() as session:
                rec = await session.run(
                    "MATCH (n) WHERE n.code = $raw OR n.name = $raw "
                    "RETURN id(n) AS id LIMIT 1", raw=raw)
                row = await rec.single()
                return row["id"] if row else None
        except Exception:
            return None

    if not body.node_ids:
        return {"ok": False, "messages": ["❌ 未指定推理节点"]}

    # 取第一个节点作为种子
    seed_id = await _resolve_seed_id(body.node_ids[0])
    if seed_id is None:
        return {"ok": False, "messages": [f"❌ 无法找到节点: {body.node_ids[0]}"]}

    result = await run_reasoning(
        seed_node_id=seed_id,
        copy_version=body.copy_version,
        confidence_threshold=body.confidence,
    )

    return {
        "ok": result["error"] is None,
        "messages": result["log"],
        "copy_version": result["copy_version"],
        "clone_count": result["clone_count"],
        "edges_built": result["edges_built"],
    }


@router.get("/ontology/search")
async def search_nodes(keyword: str, limit: int = 20, graph=Depends(get_graph)):
    """按关键词搜索节点。"""
    try:
        return await graph.search_nodes(keyword, limit=min(limit, 100))
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Graph DB search failed: {e}")


# =========================================================================
# Ontology Model CRUD (SQLite)
# =========================================================================

@router.get("/ontology-models")
async def list_ontology_models(
    keyword: Optional[str] = None,
    limit: int = 50,
    repo=Depends(get_ontology_repo),
):
    """获取本体模型树（含继承字段数）。"""
    try:
        if keyword:
            return await repo.search_models(keyword, limit=min(limit, 200))
        tree = await repo.get_full_tree_with_attrs()
        # 补全继承字段数：子模型 badge 显示完整字段数而非 0
        for node in tree:
            if not node.get("attributes"):
                inherited = _get_inherited_fields(node["id"])
                node["_inherited_count"] = len(inherited)
        return tree
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Database query failed: {e}")


@router.get("/ontology-models/stats")
async def ontology_models_stats(repo=Depends(get_ontology_repo)):
    """获取本体模型统计。"""
    try:
        return await repo.get_stats()
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Database query failed: {e}")


@router.get("/ontology-models/search")
async def search_ontology_models(
    keyword: str,
    model_type: Optional[str] = None,
    status: Optional[str] = None,
    limit: int = 50,
    repo=Depends(get_ontology_repo),
):
    """搜索本体模型。"""
    try:
        if not model_type and not status:
            return await repo.search_models(keyword, limit=min(limit, 1000))
        from infrastructure.db.base_repo import BaseRepository
        pool = get_sqlite_pool()
        temp = BaseRepository(pool, "ontol_model", pk="id", soft_delete=True)
        where = {}
        if model_type:
            where["ontol_model_type"] = model_type
        if status:
            where["ontol_model_status"] = status
        return await temp.search(keyword, columns=["name", "ontol_model_desc"], where=where, limit=limit)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Search failed: {e}")


@router.get("/ontology-models/types/flat")
async def list_ontology_types_flat():
    """获取所有本体类型扁平列表（供下拉框使用），含字段数。"""
    types = _load_ontology_types()
    result = []
    for tid, t in types.items():
        inherited = _get_inherited_fields(tid)
        result.append({
            "id": tid,
            "name": t["name"],
            "type_code": t["type_code"],
            "parent_id": t["parent_id"],
            "desc": t.get("desc", ""),
            "field_count": len(inherited),
        })
    result.sort(key=lambda x: (x["type_code"] or "", x["name"]))
    return result


@router.get("/ontology-models/{model_id}/inherited-fields")
async def get_model_inherited_fields(model_id: str):
    """获取指定本体模型的完整字段列表（含继承链：M_ROOT → ... → 当前模型）。

    子类型字段覆盖父类型同名字段，返回以 code 为键的字段列表。
    """
    from fastapi import HTTPException as _HTTPException
    types = _load_ontology_types()
    if model_id not in types:
        raise _HTTPException(status_code=404, detail=f"本体模型不存在: {model_id}")
    fields = _get_inherited_fields(model_id)
    # 按 source 分组排序（M_ROOT 的字段在前，当前模型的在后）
    sorted_fields = sorted(fields.values(), key=lambda f: (f.get("source_model", "") == model_id, f.get("order", 0), f.get("code", "")))
    return {
        "model_id": model_id,
        "model_name": types[model_id]["name"],
        "type_code": types[model_id]["type_code"],
        "field_count": len(sorted_fields),
        "fields": sorted_fields,
    }


@router.get("/ontology-models/{model_id}")
async def get_ontology_model(model_id: str, repo=Depends(get_ontology_repo)):
    """获取单个本体模型及其属性（含继承链）。"""
    try:
        if model_id != "tree":
            model = await repo.get_model_with_attrs(model_id)
        else:
            return await repo.get_full_tree_with_attrs()
        if model is None:
            model = await repo.model.get_by_id(model_id)
        return model
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Database query failed: {e}")


@router.post("/ontology-models", status_code=201)
async def create_ontology_model(body: OntolModelCreateBody, repo=Depends(get_ontology_repo)):
    """创建本体模型（后端生成 UUID 主键，code 由用户指定并唯一校验）。"""
    import uuid as _uuid
    data = body.model_dump()
    data["id"] = _uuid.uuid4().hex[:16]  # 主键自动 UUID
    try:
        return await repo.model.insert(data)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Create failed: {e}")


@router.put("/ontology-models/{model_id}")
async def update_ontology_model(model_id: str, body: OntolModelUpdateBody, repo=Depends(get_ontology_repo)):
    """更新本体模型。"""
    try:
        result = await repo.model.update(model_id, body.model_dump(exclude_none=True))
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Update failed: {e}")


@router.delete("/ontology-models/{model_id}")
async def delete_ontology_model(model_id: str, soft: bool = True, repo=Depends(get_ontology_repo)):
    """删除本体模型（默认软删除）。"""
    try:
        deleted = await repo.model.delete(model_id, soft=soft)
        return {"deleted": True, "model_id": model_id, "soft": soft}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete failed: {e}")


@router.get("/ontology-models/{model_id}/attrs")
async def list_model_attrs(model_id: str, is_system: Optional[str] = None, repo=Depends(get_ontology_repo)):
    """获取模型属性列表。"""
    try:
        return await repo.get_attrs_by_model(model_id, is_system=is_system)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Database query failed: {e}")


@router.post("/ontology-models/{model_id}/attrs", status_code=201)
async def create_model_attr(model_id: str, body: OntolModelAttrCreateBody, repo=Depends(get_ontology_repo)):
    """创建模型属性（后端生成 UUID 主键，前端传的 id 会被忽略）。"""
    import uuid as _uuid
    data = body.model_dump()
    data["id"] = _uuid.uuid4().hex[:16]          # 用 UUID 覆盖前端传的 id
    data["ontol_model_id"] = model_id
    data.setdefault("create_time", datetime.utcnow())
    data.setdefault("delete_flag", "0")
    try:
        return await repo.attr.insert(data)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Create attr failed: {e}")


@router.put("/ontology-models/{model_id}/attrs/{attr_id}")
async def update_model_attr(model_id: str, attr_id: str, body: OntolModelAttrUpdateBody, repo=Depends(get_ontology_repo)):
    """更新模型属性。系统预设字段（attr_is_system='1'）不可修改。"""
    try:
        # 检查是否为系统预设字段
        existing = await repo.attr.get_by_id(attr_id)
        if existing and existing.get("attr_is_system") == "1":
            raise HTTPException(status_code=403, detail="系统预设字段不可修改")
        result = await repo.attr.update(attr_id, body.model_dump(exclude_none=True))
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Update attr failed: {e}")


@router.delete("/ontology-models/{model_id}/attrs/{attr_id}")
async def delete_model_attr(model_id: str, attr_id: str, repo=Depends(get_ontology_repo)):
    """删除模型属性。系统预设字段（attr_is_system='1'）不可删除。"""
    try:
        # 检查是否为系统预设字段
        existing = await repo.attr.get_by_id(attr_id)
        if existing and existing.get("attr_is_system") == "1":
            raise HTTPException(status_code=403, detail="系统预设字段不可删除")
        await repo.attr.delete(attr_id)
        return {"deleted": True, "attr_id": attr_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete attr failed: {e}")


# =========================================================================
# Upload & File Management
# =========================================================================
import os
import json as _json
import time
import threading
from datetime import datetime

class SnowflakeGenerator:
    """简化版雪花算法，生成 64 位整数 ID。

    结构: timestamp(42bit) | datacenter(5bit) | worker(5bit) | sequence(12bit)
    """

    EPOCH = 1577836800000  # 2020-01-01T00:00:00Z (ms)

    def __init__(self, worker_id: int = 1, datacenter_id: int = 1):
        self.worker_id = worker_id & 0x1F
        self.datacenter_id = datacenter_id & 0x1F
        self.sequence = 0
        self.last_timestamp = -1
        self._lock = threading.Lock()

    def next_id(self) -> int:
        with self._lock:
            timestamp = int(time.time() * 1000)
            if timestamp < self.last_timestamp:
                timestamp = self.last_timestamp

            if timestamp == self.last_timestamp:
                self.sequence = (self.sequence + 1) & 0xFFF
                if self.sequence == 0:
                    while timestamp <= self.last_timestamp:
                        timestamp = int(time.time() * 1000)
            else:
                self.sequence = 0

            self.last_timestamp = timestamp
            return (
                ((timestamp - self.EPOCH) << 22)
                | (self.datacenter_id << 17)
                | (self.worker_id << 12)
                | self.sequence
            )


def _generate_snowflake_ids(
    entities: list[dict],
    relationships: list[dict],
    existing_ids: set[int],
) -> dict[str, int]:
    """为实体中 LLM 随机生成的 id 分配雪花 ID，相同随机串 → 相同雪花 ID。

    Returns:
        {random_id_str: snowflake_int} 映射表
    """
    # 收集所有需要替换的随机 ID 字符串
    random_ids: set[str] = set()
    for ent in entities:
        eid = (ent.get("properties", {}).get("id") or "").strip()
        if not eid:
            continue
        # 跳过纯数字（可能已是雪花 ID）和提示词占位文本
        if eid.isdigit() or eid == "大模型随机生成":
            continue
        random_ids.add(eid)

    # 也检查关系的 start_node_id / end_node_id 是否引用了随机 ID
    for rel in relationships:
        for key in ("start_node_id", "end_node_id", "subject", "object"):
            val = (rel.get(key) or "").strip()
            if val and not val.isdigit():
                random_ids.add(val)

    if not random_ids:
        return {}

    # 为每个唯一的随机 ID 生成雪花 ID，确保不与已有 ID 冲突
    sf = SnowflakeGenerator(worker_id=1, datacenter_id=1)
    id_map: dict[str, int] = {}
    for rid in random_ids:
        while True:
            snowflake = sf.next_id()
            if snowflake not in existing_ids and snowflake not in id_map.values():
                id_map[rid] = snowflake
                existing_ids.add(snowflake)
                break

    return id_map


_UPLOAD_HISTORY_FILE = _Path("infrastructure/storage/uploads/.history.json")


def _read_history() -> list:
    if _UPLOAD_HISTORY_FILE.exists():
        try:
            return _json.loads(_UPLOAD_HISTORY_FILE.read_text(encoding="utf-8"))
        except Exception:
            return []
    return []


def _write_history(entries: list):
    _UPLOAD_HISTORY_FILE.parent.mkdir(parents=True, exist_ok=True)
    _UPLOAD_HISTORY_FILE.write_text(_json.dumps(entries, ensure_ascii=False, indent=2), encoding="utf-8")


def _append_history(entry: dict):
    entries = _read_history()
    entries.insert(0, entry)
    # keep last 200
    _write_history(entries[:200])


@router.get("/upload/history")
async def upload_history():
    """返回历史上传记录。"""
    return _read_history()


@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """接收文件上传，保存到 infrastructure/storage/ 目录。"""
    from pathlib import Path as _PathLocal

    upload_dir = _PathLocal("infrastructure/storage/uploads")
    upload_dir.mkdir(parents=True, exist_ok=True)

    safe_name = os.path.basename(file.filename or "untitled")
    dest = upload_dir / safe_name

    try:
        content = await file.read()
        if len(content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File exceeds 50MB limit")

        dest.write_bytes(content)
        now_str = datetime.utcnow().isoformat() + "Z"

        entry = {
            "filename": safe_name,
            "size": len(content),
            "path": str(dest),
            "uploaded_at": now_str,
            "uploaded": True,
        }
        _append_history(entry)

        return entry
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/upload/preview/{filename:path}")
async def preview_file(filename: str):
    """在线预览上传的文件。文本文件返回内容，二进制文件返回 download URL。"""
    from fastapi.responses import FileResponse, PlainTextResponse

    upload_dir = _Path("infrastructure/storage/uploads")
    safe_name = os.path.basename(filename)
    file_path = upload_dir / safe_name

    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"File '{safe_name}' not found")

    # 文本类直接返回内容
    text_exts = {".txt", ".json", ".csv", ".xml", ".rdf", ".owl", ".ttl", ".nt", ".swrl"}
    if file_path.suffix.lower() in text_exts:
        try:
            content = file_path.read_text(encoding="utf-8")
            return {"type": "text", "filename": safe_name, "content": content}
        except UnicodeDecodeError:
            pass

    return FileResponse(str(file_path), filename=safe_name)


@router.delete("/upload/{filename:path}")
async def delete_file(filename: str):
    """删除上传的文件（同时清理历史记录）。"""
    upload_dir = _Path("infrastructure/storage/uploads")
    safe_name = os.path.basename(filename)
    file_path = upload_dir / safe_name

    file_deleted = False
    if file_path.exists():
        try:
            file_path.unlink()
            file_deleted = True
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    # 无论文件是否存在，都从历史记录中移除该条目
    history = _read_history()
    remaining = [e for e in history if e.get("filename") != safe_name]
    if len(remaining) != len(history):
        _write_history(remaining)

    return {"deleted": True, "filename": safe_name, "file_on_disk": file_deleted}


# =========================================================================
# 本体类型感知的实体解析 & 图数据库导入
# =========================================================================

# ── 本体类型定义（从 SQLite 动态加载）──

_ONTOLOGY_TYPES_CACHE: Optional[dict] = None
_ONTOLOGY_TREE_CACHE: Optional[list] = None   # 树形结构（含层级关系）


def _load_ontology_types() -> dict:
    """从 SQLite 数据库加载所有本体类型及其字段定义（平铺）。"""
    global _ONTOLOGY_TYPES_CACHE
    if _ONTOLOGY_TYPES_CACHE is not None:
        return _ONTOLOGY_TYPES_CACHE

    import sqlite3
    from pathlib import Path

    db_path = Path("infrastructure/db/ontol.db")
    types = {}

    if db_path.exists():
        conn = sqlite3.connect(str(db_path))
        conn.row_factory = sqlite3.Row
        models = conn.execute(
            "SELECT * FROM ontol_model WHERE delete_flag='0' ORDER BY ontol_model_type, id"
        ).fetchall()
        for m in models:
            md = dict(m)
            attrs = conn.execute(
                """SELECT id, name, code, attr_data_type, attr_length,
                          attr_required, attr_default_value, attr_desc, attr_order
                   FROM ontol_model_attr
                   WHERE ontol_model_id=? AND delete_flag='0'
                   ORDER BY attr_order, code""",
                (md["id"],),
            ).fetchall()
            types[md["id"]] = {
                "id": md["id"],
                "name": md["name"],
                "parent_id": md.get("ontol_parent_id") or None,
                "type_code": md["ontol_model_type"],
                "desc": md["ontol_model_desc"] or "",
                "fields": [
                    {
                        "id": a["id"],
                        "name": a["name"],
                        "code": a["code"],
                        "data_type": a["attr_data_type"],
                        "length": a["attr_length"],
                        "required": a["attr_required"],
                        "default": a["attr_default_value"],
                        "desc": a["attr_desc"] or "",
                        "order": a["attr_order"],
                    }
                    for a in attrs
                ],
            }
        conn.close()

    _ONTOLOGY_TYPES_CACHE = types
    return types


def _get_inherited_fields(ont_type: str) -> dict[str, dict]:
    """
    获取指定本体类型的完整字段列表（含继承链）。

    规则：
    - M_ROOT 的字段是所有类型共用的（L0 全局字段）
    - 沿父链向上递归收集，直到 M_ROOT
    - 子类型字段覆盖父类型同名字段
    - 返回 {field_code: field_info} 字典
    """
    types = _load_ontology_types()
    if not types or ont_type not in types:
        return {}

    fields: dict[str, dict] = {}
    current_id = ont_type
    visited = set()

    while current_id and current_id in types and current_id not in visited:
        visited.add(current_id)
        t = types[current_id]
        for f in t.get("fields", []):
            code = f.get("code", "")
            if code not in fields:
                fields[code] = {**f, "source_model": current_id, "source_name": t.get("name", "")}
        current_id = t.get("parent_id")

    return fields


def _validate_entities_for_import(entities: list[dict]) -> dict:
    """
    校验解析后的实体，检查本体模板匹配 + 计算缺失字段。

    返回:
        {
            "valid": bool,
            "type_counts": {ont_type: count},
            "unknown_types": [{ont_type, type_name, entity_names: [...]}],
            "missing_fields": [{entity_name, ont_type, missing: [{code, name, required, default}]}],
            "summary": "校验摘要",
        }
    """
    types = _load_ontology_types()
    all_type_ids = set(types.keys()) if types else set()

    type_counts: dict[str, int] = {}
    unknown_types_map: dict[str, dict] = {}  # ont_type -> {ont_type, type_name, entity_names}
    missing_fields_list: list[dict] = []

    for ent in entities:
        ont_type = (ent.get("ont_type") or "").strip()
        name = (ent.get("name") or "").strip()
        type_name = (ent.get("type_name") or "").strip()
        props = ent.get("properties") or {}

        # 统计
        type_counts[ont_type] = type_counts.get(ont_type, 0) + 1

        # 检查模板是否存在
        if ont_type not in all_type_ids:
            key = ont_type
            if key not in unknown_types_map:
                unknown_types_map[key] = {"ont_type": ont_type, "type_name": type_name, "entity_names": []}
            if name and name not in unknown_types_map[key]["entity_names"]:
                unknown_types_map[key]["entity_names"].append(name)
            continue

        # 模板存在 → 计算缺失字段
        inherited = _get_inherited_fields(ont_type)
        missing = []
        for code, finfo in inherited.items():
            if code not in props or not props.get(code):
                missing.append({
                    "code": code,
                    "name": finfo.get("name", ""),
                    "required": finfo.get("required", "0") == "1",
                    "default": finfo.get("default") or "",
                    "source_model": finfo.get("source_model", ""),
                    "source_name": finfo.get("source_name", ""),
                })
        if missing:
            missing_fields_list.append({
                "entity_name": name,
                "ont_type": ont_type,
                "type_name": type_name,
                "missing": missing,
                "has_existing_props": list(props.keys()),
            })

    unknown_types = list(unknown_types_map.values())
    has_unknown = len(unknown_types) > 0
    has_missing = len(missing_fields_list) > 0
    valid = not has_unknown

    summary_parts = []
    if has_unknown:
        unknown_names = ", ".join(f"{u['ont_type']}({len(u['entity_names'])}个实体)" for u in unknown_types[:5])
        summary_parts.append(f"⚠️ {len(unknown_types)} 个类型无匹配模板: {unknown_names}")
    if has_missing:
        summary_parts.append(f"📋 {len(missing_fields_list)} 个实体可补全字段")
    if not summary_parts:
        summary_parts.append("✅ 所有实体类型均有匹配模板")

    return {
        "valid": valid,
        "type_counts": type_counts,
        "unknown_types": unknown_types,
        "missing_fields": missing_fields_list,
        "summary": " | ".join(summary_parts),
    }


def _build_ontology_prompt() -> str:
    """构建包含所有本体类型定义的 LLM 提示词。

    使用 _get_inherited_fields() 获取每个类型的完整继承字段（M_ROOT → 父链 → 自身），
    确保 AI 看到的字段定义与 API 返回的继承字段一致。
    """
    types = _load_ontology_types()
    root_type = types.get("M_ROOT", {})
    non_root = {tid: td for tid, td in types.items() if tid != "M_ROOT"}

    lines = []
    lines.append("你是一个本体建模专家。请解析文本，识别实体并归类到以下本体类型，填写所有字段。\n")

    # ─── 1. M_ROOT 根节点概览（所有类型的字段基础）───
    root_inherited = _get_inherited_fields("M_ROOT") if "M_ROOT" in types else {}
    if root_inherited:
        lines.append(f"# M_ROOT 本体根节点 — {len(root_inherited)} 个基础字段（所有类型继承）\n")
        for f in sorted(root_inherited.values(), key=lambda x: x.get("code", "")):
            lines.append(_fmt_field(f))
        lines.append("")

    # ─── 2. 各本体类型 + 完整继承字段 ──
    lines.append("# 本体类型定义\n")
    for tid, tdef in non_root.items():
        lines.append(f"## {tdef.get('name','')} (ont_type={tid}, 类型代码={tdef.get('type_code','')})")
        lines.append(f"描述: {tdef.get('desc','')}")

        # 用 _get_inherited_fields 获取完整字段（M_ROOT + 父链 + 自身）
        all_fields = _get_inherited_fields(tid)
        own_fields = [f for f in all_fields.values() if f.get("source_model") == tid]
        inherited_from = [f for f in all_fields.values() if f.get("source_model") != tid]

        lines.append(f"完整字段 ({len(all_fields)} 个归属, {len(inherited_from)} 个继承 + {len(own_fields)} 个归属):")
        for f in sorted(all_fields.values(), key=lambda x: (x.get("source_model", "") == tid, x.get("order", 0), x.get("code", ""))):
            src_tag = "" if f.get("source_model") == tid else f" [继承自 {f.get('source_name','')}]"
            lines.append(_fmt_field(f) + src_tag)
        lines.append("")

    # ─── 3. 字段汇总表 ──
    lines.append("# 字段汇总\n")
    lines.append("每个实体需填写的字段 = 该本体类型的完整继承字段列表（含 M_ROOT 基础字段）。\n")
    for tid, tdef in non_root.items():
        all_fields = _get_inherited_fields(tid)
        codes = [f.get("code", "") for f in sorted(all_fields.values(), key=lambda x: x.get("code", ""))]
        lines.append(f"{tdef['name']} ({tid}, 共{len(codes)}字段): {', '.join(codes)}")
    lines.append("")

    # ─── 4. 本体类型枚举说明 ──
    lines.append("""# 本体类型（Ontology Types）枚举说明

在解析文本时，请根据实体的核心业务特征，将其归类为以下 7 种本体类型之一。注意：输出 JSON 时，`type` 字段的值必须严格使用以下指定的代码（如 M1、M2 等）：
- **M1 实体 (M_ENTITY)**：描述客观存在的物理对象、数字资产或核心业务概念。例如：设备、传感器、产品、文档、数据表等。
- **M2 行为 (M_ACTION)**：描述主体执行的动作、操作、流程节点或状态变更。例如：启动、校验、清洗、审批、计算等。
- **M3 规则 (M_RULE)**：描述业务逻辑、约束条件、算法策略、触发条件或计算公式。例如：阈值告警规则、权限校验规则、调度策略等。
- **M4 场景 (M_SCENE)**：描述业务发生的上下文、环境、时间段或特定业务模式。例如：夜间巡检、高并发交易、设备离线状态等。
- **M5 主体 (M_SUBJECT)**：描述执行行为的发起者、责任方、参与角色或组织。例如：操作员、系统服务、部门、外部供应商等。
- **M6 异常 (M_EXCEPTION)**：描述偏离正常状态的故障、错误、风险或告警事件。例如：网络超时、数据缺失、设备过热、越权访问等。
- **M7 质量 (M_QUALITY)**：描述衡量业务、数据或系统表现的标准、指标或评估维度。例如：准确率、响应时间、完整性、合规性等。

""")

    # ─── 5. 动态生成输出示例 JSON（用第一个非 root 类型的字段）───
    example_type = list(non_root.keys())[0] if non_root else None
    example_fields = _get_inherited_fields(example_type) if example_type else {}
    example_props_lines = []
    if example_fields:
        for f in sorted(example_fields.values(), key=lambda x: x.get("code", "")):
            code = f.get("code", "")
            name = f.get("name", "")
            dtype = f.get("data_type", "VARCHAR")
            default = f.get("default", "")
            desc = f.get("desc", "")
            dtype_hint = dtype
            val_hint = default if default else (f"示例{name}" if dtype == "0" else "0")
            example_props_lines.append(f'        "{code}": "{val_hint}",  // {name}, {dtype_hint}{", " + desc if desc else ""}')
    example_props_json = "\n".join(example_props_lines) if example_props_lines else '        // (无定义字段)'

    # ─── 6. 语义规范: OWL2 DL + SWRL + SHACL ──
    lines.append("""# 语义规范

## 前缀约定

| 序号 | 作用域 | 名称 | 编码前缀 | 格式示例 | 备注 |
|------|--------|------|----------|----------|------|
| 1 | 对象属性 | RDFS语言 | `rdfs:` | 也支持RDFS核心常量，不写前缀 | RDFS语言 |
| 2 | 对象属性 | OWL2 DL语言 | `owl2:` | | OWL2 DL语言为主 |
| 3 | 对象属性 | SWRL语言 | `swrl:` | | SWRL语法 |
| 4 | 对象属性 | SHACL语言 | `sh:` | | |
| 5 | 对象属性 | 规则设定 | `rule:` | `rule:forwardChain`<br>`rule:backwardChain` | 默认就是前链推理 |
| 6 | 对象属性 | 自定义动态函数 | `func:` | `{"id":"图ID","func":"函数名"}` | 不对接大模型，用JSON调用函数实现 |
| 7 | 边属性 | 自定义动作接口 | 边的属性 | `actionType: "inference"`<br>`required: true`<br>`validationType: "Strong"`<br>`ruleId: "phone_format_rule_001"`<br>`func: "validate_phone_format"`<br>`id: "field_B_node_123"`<br>`msg: "详细说明作用"`<br>`synonym: "同义词"`<br>`queryVariant: "错意词"` | actionType (路由标识)：指定执行分支（如 inference 表示走推理机逻辑）。<br>required (阻断控制)：定义校验失败时，是否强制中断当前业务流程。<br>validationType (规则级别)：声明校验的严格程度（如 Strong 强校验（阻断），弱校验Weak（提醒不阻断））。<br>ruleId (规则锚点)：指向图数据库中的"规则本体节点"，用于元数据管理和错误信息溯源。<br>func (执行指令)：直接映射底层要调用的具体函数名，保障执行引擎高效运转。<br>id (数据锚点)：明确当前需要被校验的具体业务数据节点。 |

> ⚠️ **核心约束**：Memgraph 边属性仅支持标量类型（String / Int / Float / Bool / DateTime / Duration / Point / List）。不支持 Map / JSON 嵌套。所有复合语义必须通过「扁平化 key-value」或「独立节点 + 关系」表达。

## RDFS 核心词汇
RDFS（RDF Schema）为 RDF 数据模型提供基础的类型系统和词汇描述能力：
- **rdfs:subClassOf** — 子类关系（A 是 B 的子类型）
- **rdfs:subPropertyOf** — 子属性关系
- **rdfs:domain** — 属性定义域（该属性适用于哪类主体）
- **rdfs:range** — 属性值域（该属性的值属于哪类客体）
- **rdfs:label** — 人类可读标签
- **rdfs:comment** — 注释说明
- **rdfs:type** — 实例类型声明
- **rdfs:Class** — 类
- **rdfs:Property** — 属性

## 关系（predicate）— 使用 OWL2 DL 语义
实体之间的关系遵循 OWL2 DL（Description Logic）标准，所有关系类型使用 `owl2:` 前缀：
- **owl2:subClassOf** — 子类关系（A 是 B 的子类型）
- **owl2:equivalentClass** — 等价类关系
- **owl2:disjointWith** — 互斥关系（A 和 B 不能同时成立）
- **owl2:objectProperty** — 对象属性（A 指向 B 的语义关联）
- **owl2:dataProperty** — 数据属性（A 拥有某个数据值）
- **owl2:sameAs / owl2:differentFrom** — 个体等价/不等价
- **owl2:inverseOf** — 逆关系（A→B 和 B→A 互为逆）
- **owl2:domain / owl2:range** — 定义域的约束

关系 type 优先使用 OWL2 标准词汇（带 owl2: 前缀），如需要自定义，使用驼峰命名（如 hasPart、isLocatedAt）。

## 推理规则（SWRL）
如文本中包含推理规则，用 SWRL（Semantic Web Rule Language）表达，示例：
  - swrl:Antecedent(body) → swrl:Consequent(head)
  - Entity(?x) ^ hasProperty(?x, ?v) ^ swrlb:greaterThan(?v, 100) → HighValue(?x)

## 校验规则（SHACL）
如文本中包含数据校验约束，用 SHACL（Shapes Constraint Language）表达，常用词汇：
  - **sh:property** — 属性约束声明
  - **sh:class** — 节点类型约束
  - **sh:datatype** — 数据类型约束（如 xsd:string、xsd:integer）
  - **sh:minCount / sh:maxCount** — 最小/最大出现次数
  - **sh:pattern** — 正则表达式匹配
  - **sh:in** — 枚举值约束

## 推理规则设定（rule:）
如文本中包含推理方向或策略设定，使用 `rule:` 前缀表达：
  - `rule:forwardChain` — 前链推理（从已知事实推导新结论），**默认模式**
  - `rule:backwardChain` — 后链推理（从目标反向寻找支撑条件）

## 自定义动态函数（func:）
如文本中需要执行自定义计算或处理逻辑，使用 `func:` 前缀表达。不对接大模型，直接通过 JSON 调用底层函数实现：
  - 格式：`{"id": "图节点ID", "func": "函数名"}`
  - 函数参数根据具体业务需求扩展

## 自定义动作接口（边属性）
边属性用于在关系边上附加校验和执行控制信息，使用 Memgraph key-value 形式存储，不支持嵌套 JSON：

| 字段 | 类型 | 说明 |
|------|------|------|
| `actionType` | string | 路由标识，指定执行分支（如 `inference` 表示走推理机逻辑） |
| `required` | boolean | 阻断控制，校验失败时是否强制中断当前业务流程 |
| `validationType` | string | 规则级别，`Strong` 强校验（阻断），`Weak` 弱校验（提醒不阻断） |
| `ruleId` | string | 规则锚点，指向图数据库中的规则本体节点，用于元数据管理和错误信息溯源 |
| `func` | string | 执行指令，直接映射底层要调用的具体函数名，保障执行引擎高效运转 |
| `id` | string | 数据锚点，明确当前需要被校验的具体业务数据节点 |
| `msg` | string | 详细说明该边属性的作用 |
| `synonym` | string | 同义词，用于语义匹配和模糊查询 |
| `queryVariant` | string | 错意词/变体词，用于容错查询 |

示例：
```
actionType: "inference"
required: true
validationType: "Strong"
ruleId: "phone_format_rule_001"
func: "validate_phone_format"
id: "field_B_node_123"
msg: "校验电话号码格式是否符合规范"
synonym: "手机号校验"
queryVariant: "电话验证,号码检查"
```

## 字段填写规则
1. 每个字段尽量从原文中推断填充，无法推断则留空字符串 ""
2. 字段值保持原文语义，不要臆造
3. 日期/时间字段使用 openCypher 标准时间格式（如 LocalDateTime），到秒即可
4. 编码字段使用英文驼峰或下划线命名
5. 置信度字段如未明确指定，默认填 80%

## 输出格式（严格 JSON）
只输出以下 JSON，不得包含任何解释文字：

```json
{
  "entities": [
    {
      "name": "实体名称",
      "ont_type": "M_ENTITY",
      "type_name": "实体",
      "properties": {
""")
    lines.append(example_props_json)
    lines.append("""      }
    }
  ],
  "relationships": [
    {
      "type": "关系类型(如 owl2:subClassOf)",
      "start_node_id": "起始节点名称",
      "end_node_id": "目标节点名称",
      "properties": {
        "note": "Memgraph 属性图模型支持在关系上挂载属性，按需填写"
      }
    }
  ]
}
```

## 重要规则
1. 每个实体必须归类到一个 ont_type（从上面定义的本体类型中选择）
2. 每个实体的 properties 中 `type` 字段必须填写 M1~M7 枚举值
3. properties 必须包含该类型的完整继承字段列表，尽量从文本中提取
4. 关系 type 遵循 OWL2 DL 语义规范（带 owl2: 前缀）
5. 如有推理规则，使用 SWRL 格式填入 hasPrecondition 或单独关系
6. 如有校验约束，使用 SHACL 格式
7. 只输出 JSON，不要输出任何解释""")
    return "\n".join(lines)


def _fmt_field(f: dict) -> str:
    """格式化一个字段为提示词中的一行。"""
    req = "必填" if f.get("required", "0") == "1" else "可选"
    default = f"，默认值={f.get('default')}" if f.get("default") else ""
    dtype_name = f.get("data_type", "VARCHAR")
    return f"  - {f.get('code','')} ({f.get('name','')}): {req}, {dtype_name}, 长度={f.get('length','—')}{default} — {f.get('desc','')}"


class ParseTriplesRequest(BaseModel):
    filename: str = Field(..., description="要解析的文件名")
    model: str = Field("", description="使用的 LLM 模型名 (DB config_id 或 provider::name)")


class ImportEntitiesRequest(BaseModel):
    filename: str = Field(..., description="来源文件名")
    entities: list[dict] = Field(..., description="本体实体列表")
    relationships: list[dict] = Field(default_factory=list, description="关系列表")
    scene_ids: list[str] = Field(default_factory=list, description="关联的场景ID列表")


class ValidateEntitiesRequest(BaseModel):
    entities: list[dict] = Field(..., description="待校验的实体列表")


@router.post("/upload/validate-entities")
async def validate_entities_for_import(body: ValidateEntitiesRequest):
    """
    校验解析后的实体，检测本体模板匹配情况。

    返回:
    - unknown_types: 无匹配模板的类型及其实体（需用户确认）
    - missing_fields: 有模板但缺失字段的实体（需补全）
    """
    from pathlib import Path as _PathLocal
    return _validate_entities_for_import(body.entities)


def _parse_entities_json(text: str) -> dict:
    """从 LLM 输出中提取 JSON 格式的实体和关系（多级降级策略）。"""
    import re, json

    text = text.strip()

    # 1. 直接 JSON 解析
    try:
        result = json.loads(text)
        if isinstance(result, dict):
            return result
    except (json.JSONDecodeError, TypeError):
        pass

    # 2. 提取 markdown 代码围栏（各种变体）
    fence_patterns = [
        r'```json\s*\n(.*?)\n\s*```',
        r'```\s*\n(.*?)\n\s*```',
        r'```json\s*(.*?)\s*```',
        r'```\s*(.*?)\s*```',
    ]
    for pattern in fence_patterns:
        for m in re.finditer(pattern, text, re.DOTALL | re.IGNORECASE):
            try:
                result = json.loads(m.group(1).strip())
                if isinstance(result, dict):
                    return result
            except (json.JSONDecodeError, TypeError):
                continue

    # 3. 按大括号平衡匹配，逐个尝试每个顶层 JSON 对象
    brace_depth = 0
    start = -1
    candidates = []
    for i, ch in enumerate(text):
        if ch == '{':
            if brace_depth == 0:
                start = i
            brace_depth += 1
        elif ch == '}':
            brace_depth -= 1
            if brace_depth == 0 and start >= 0:
                candidates.append(text[start:i+1])
                start = -1

    for candidate in candidates:
        try:
            result = json.loads(candidate)
            if isinstance(result, dict) and ('entities' in result or 'relationships' in result):
                return result
        except (json.JSONDecodeError, TypeError):
            continue

    # 4. 最终降级：处理常见 LLM 附加文本后重试贪婪匹配
    cleaned = re.sub(r'\*\*[^*]+\*\*', '', text)       # 去加粗
    cleaned = re.sub(r'`[^`]+`', '', cleaned)           # 去行内代码
    m = re.search(r'\{.*\}', cleaned, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(0))
        except (json.JSONDecodeError, TypeError):
            pass

    # 全失败：返回空结果 + 错误上下文
    return {
        "entities": [],
        "relationships": [],
        "_parse_error": True,
        "_raw_snippet": text[:500],
    }


def _build_classify_prompt() -> str:
    """构建分类提示词 — 仅让 LLM 判断实体归类，不涉及字段细节。"""
    types = _load_ontology_types()
    non_root = [(tid, td) for tid, td in types.items() if tid != "M_ROOT"]

    lines = [
        "你是一个本体建模专家。请解析文本，识别实体并归类到以下本体类型。",
        "只需返回实体名称和本体类型，不需要填写属性字段。\n",
        "# 本体类型枚举\n",
    ]
    for tid, tdef in non_root:
        lines.append(f"- **{tdef.get('name','')}** (ont_type={tid}, type_code={tdef.get('type_code','')}): {tdef.get('desc','')}")

    lines.append("""
# 输出格式（严格 JSON，只输出此 JSON）
```json
{
  "entities": [
    {"name": "实体名称", "ont_type": "M_ENTITY"}
  ],
  "relationships": [
    {"type": "关系类型", "start_node_id": "起始实体名", "end_node_id": "目标实体名"}
  ]
}
```""")
    return "\n".join(lines)


def _build_extract_prompt(ont_type: str) -> str:
    """构建字段提取提示词 — 仅针对单个本体类型，用 _get_inherited_fields 取完整字段。"""
    t = _load_ontology_types().get(ont_type, {})
    fields = _get_inherited_fields(ont_type)
    type_name = t.get("name", ont_type)
    type_code = t.get("type_code", "")

    lines = [
        f"你是本体建模专家。请为以下实体提取字段值，该实体类型为 **{type_name}** (ont_type={ont_type}, type_code={type_code})。",
        f"\n# 需填写的字段（共 {len(fields)} 个，含继承字段）\n",
    ]
    for f in sorted(fields.values(), key=lambda x: (x.get("source_model", "") == ont_type, x.get("order", 0), x.get("code", ""))):
        src = f" [继承自 {f['source_name']}]" if f.get("source_model") != ont_type else ""
        lines.append(_fmt_field(f) + src)

    lines.append("""
# 输出格式（严格 JSON，只输出此 JSON）
```json
{
  "entities": [
    {
      "name": "实体名称（与输入一致）",
      "ont_type": \"""" + ont_type + """\",
      "properties": {
        "field_code": "从原文提取的值，无法提取则留空"
      }
    }
  ]
}
```

## 规则
1. 每个字段尽量从原文推断，无法提取留空字符串 ""
2. 字段值保持原文语义，不要臆造
3. 置信度字段默认填 "80%"
4. 只输出 JSON，不输出解释""")
    return "\n".join(lines)


# =====================================================================
# Word 文档文本提取（用于 /upload/parse）
# =====================================================================


def _extract_text_from_docx(path: str) -> str:
    """从 .docx 文件提取纯文本。"""
    from docx import Document

    doc = Document(path)
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    content = "\n\n".join(paragraphs)
    if not content.strip():
        raise ValueError("未从 .docx 文件中提取到文本内容")
    return content


def _extract_text_from_doc(path: str) -> str:
    """从 .doc 文件提取纯文本（依赖 antiword 命令行工具）。"""
    import subprocess
    import shutil

    antiword = shutil.which("antiword")
    if not antiword:
        raise RuntimeError("服务器未安装 antiword，无法解析 .doc 文件。请将文件另存为 .docx 或 .txt 格式后重新上传。")

    result = subprocess.run(
        [antiword, path],
        capture_output=True,
        text=True,
        timeout=30,
    )
    if result.returncode != 0:
        stderr = result.stderr.strip()
        raise RuntimeError(f"antiword 解析 .doc 文件失败: {stderr or '未知错误'}")

    content = result.stdout.strip()
    if not content:
        raise ValueError("未从 .doc 文件中提取到文本内容")
    return content


@router.post("/upload/parse")
async def parse_file_to_entities(body: ParseTriplesRequest):
    """
    两阶段 AI 解析：先用分类提示词让 LLM 判定实体类型，
    再按类型 ID 通过 _get_inherited_fields 取完整字段，让 LLM 提取字段值。

    返回按本体类型分类的实体列表和关系列表，供用户审核后导入图数据库。
    """
    upload_dir = _Path("infrastructure/storage/uploads")
    safe_name = os.path.basename(body.filename)
    file_path = upload_dir / safe_name

    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"File '{safe_name}' not found")

    # ── 文本提取：支持 .txt / .docx / .doc ──
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".docx":
            content = _extract_text_from_docx(str(file_path))
        elif suffix == ".doc":
            content = _extract_text_from_doc(str(file_path))
        else:
            content = file_path.read_text(encoding="utf-8")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件解析失败: {str(e)}")

    from capabilities.models.resolver import resolve_llm
    from langchain_core.messages import HumanMessage

    llm_iface = resolve_llm(body.model)
    llm = await llm_iface.get_llm()

    # ── 分块 ──
    CHUNK_SIZE = 3000
    chunks: list[str] = []
    start = 0
    while start < len(content):
        end = min(start + CHUNK_SIZE, len(content))
        if end < len(content):
            for sep in ("\n\n", "\n", "。", "；", ". "):
                pos = content.rfind(sep, start, end)
                if pos > start + 500:
                    end = pos + len(sep)
                    break
        chunks.append(content[start:end].strip())
        start = end

    valid_chunks = [c for c in chunks if c and len(c) >= 10]
    chunk_errors: list[dict] = []

    # =====================================================================
    # 阶段 1: 分类 — LLM 判定每个实体的本体类型（轻量提示词，不含字段）
    # =====================================================================
    classify_prompt = _build_classify_prompt()
    all_classified: dict[str, dict] = {}  # name -> {name, ont_type}
    all_relationships: list[dict] = []
    seen_rels = set()

    for i, chunk in enumerate(chunks):
        if not chunk or len(chunk) < 10:
            continue
        try:
            response = await llm.ainvoke([
                HumanMessage(content=classify_prompt),
                HumanMessage(content=f"请解析以下文本，识别实体和关系（第{i+1}/{len(chunks)}块）：\n\n{chunk}"),
            ])
            text = response.content if hasattr(response, "content") else str(response)
            result = _parse_entities_json(text)
            if result.pop("_parse_error", None):
                chunk_errors.append({"chunk_index": i+1, "reason": "分类阶段 JSON 解析失败", "raw_snippet": result.pop("_raw_snippet", "")[:200]})
                continue
            for ent in result.get("entities", []):
                name = (ent.get("name") or "").strip()
                if not name:
                    continue
                if name not in all_classified:
                    all_classified[name] = {"name": name, "ont_type": ent.get("ont_type", "M_ENTITY")}
            for rel in result.get("relationships", []):
                s = (rel.get("start_node_id") or rel.get("subject") or "").strip()
                p = (rel.get("type") or rel.get("predicate") or "").strip()
                o = (rel.get("end_node_id") or rel.get("object") or "").strip()
                if not s or not o or not p:
                    continue
                key = f"{s}|{p}|{o}"
                if key not in seen_rels:
                    seen_rels.add(key)
                    all_relationships.append({"start_node_id": s, "type": p, "end_node_id": o, "properties": rel.get("properties", {})})
        except Exception as e:
            chunk_errors.append({"chunk_index": i+1, "reason": f"分类阶段 LLM 调用失败: {str(e)}"})

    if not all_classified:
        return {"filename": safe_name, "entity_count": 0, "relationship_count": len(all_relationships), "type_counts": {}, "entities": [], "relationships": all_relationships, "phase1_classified": 0, "chunk_errors": chunk_errors[:10]}

    # =====================================================================
    # 阶段 2: 字段提取 — 按类型分组，调用 _get_inherited_fields 取字段
    #           然后让 LLM 按这些字段从原文中提取值
    # =====================================================================
    by_type: dict[str, list[str]] = {}
    for name, info in all_classified.items():
        t = info["ont_type"]
        by_type.setdefault(t, []).append(name)

    all_entities: dict[str, dict] = {}
    extract_errors: list[dict] = []

    for ont_type, names in by_type.items():
        # 通过通用接口获取该类型的完整继承字段
        inherited = _get_inherited_fields(ont_type)
        if not inherited:
            for name in names:
                all_entities[name] = {"name": name, "ont_type": ont_type, "type_name": "", "properties": {}}
            continue

        extract_prompt = _build_extract_prompt(ont_type)
        entity_list_text = "\n".join(f"- {n}" for n in names)

        try:
            response = await llm.ainvoke([
                HumanMessage(content=extract_prompt),
                HumanMessage(content=f"以下实体需要提取字段值（类型={ont_type}）：\n{entity_list_text}\n\n原始文本已在上文中提供，请为每个实体提取字段值，返回 JSON："),
            ])
            text = response.content if hasattr(response, "content") else str(response)
            result = _parse_entities_json(text)
            if result.pop("_parse_error", None):
                extract_errors.append({"ont_type": ont_type, "reason": "字段提取 JSON 解析失败"})
                for name in names:
                    all_entities[name] = {"name": name, "ont_type": ont_type, "type_name": "", "properties": {}}
                continue
            for ent in result.get("entities", []):
                nm = (ent.get("name") or "").strip()
                if nm in names or nm in all_classified:
                    all_entities[nm] = {"name": nm, "ont_type": ont_type, "type_name": ent.get("type_name", ""), "properties": ent.get("properties", {})}
            for name in names:
                if name not in all_entities:
                    all_entities[name] = {"name": name, "ont_type": ont_type, "type_name": "", "properties": {}}
        except Exception as e:
            extract_errors.append({"ont_type": ont_type, "reason": f"字段提取 LLM 调用失败: {str(e)}"})
            for name in names:
                all_entities[name] = {"name": name, "ont_type": ont_type, "type_name": "", "properties": {}}

    entities_list = list(all_entities.values())

    # ── 按本体类型分组统计 ──
    type_counts = {}
    for e in entities_list:
        t = e.get("ont_type", "M_ENTITY")
        type_counts[t] = type_counts.get(t, 0) + 1

    return {
        "filename": safe_name,
        "entity_count": len(entities_list),
        "relationship_count": len(all_relationships),
        "type_counts": type_counts,
        "entities": entities_list,
        "relationships": all_relationships,
        "chunks_total": len(valid_chunks),
        "chunks_ok": len(valid_chunks) - len(chunk_errors),
        "chunks_failed": len(chunk_errors),
        "chunk_errors": chunk_errors[:10],
        "extract_errors": extract_errors[:10],
    }


@router.post("/upload/import-entities")
async def import_entities_to_neo4j(body: ImportEntitiesRequest):
    """
    将本体分类后的实体和关系导入图数据库（Memgraph/Neo4j）。
    每个实体按 ont_type 打标签，自动补全继承链上的缺失字段。
    """
    from infrastructure.db.neo4j import get_driver

    try:
        driver = await get_driver()
    except InfrastructureError:
        raise HTTPException(status_code=503, detail="Graph DB driver not initialized")

    nodes_created = 0
    edges_created = 0
    filled_fields_count = 0  # 记录补全的字段数

    async with driver.session() as session:
        # ── 查询图中已有的所有 ID（节点 + 边），用于雪花 ID 去重 ──
        existing_ids: set[int] = set()
        try:
            node_result = await session.run("MATCH (n) WHERE n.id IS NOT NULL RETURN DISTINCT n.id AS nid")
            async for rec in node_result:
                val = rec.get("nid")
                if val is not None:
                    try:
                        existing_ids.add(int(val))
                    except (ValueError, TypeError):
                        pass
            edge_result = await session.run(
                "MATCH ()-[r]->() WHERE r.id IS NOT NULL RETURN DISTINCT r.id AS rid"
            )
            async for rec in edge_result:
                val = rec.get("rid")
                if val is not None:
                    try:
                        existing_ids.add(int(val))
                    except (ValueError, TypeError):
                        pass
        except Exception:
            pass  # 查询失败不阻断导入

        # ── 为 LLM 随机生成的 id 分配雪花 ID ──
        id_map = _generate_snowflake_ids(body.entities, body.relationships, existing_ids)

        # ── 替换实体 properties.id（纯数字 int64，不转字符串）──
        for ent in body.entities:
            props = ent.get("properties") or {}
            eid = (props.get("id") or "").strip()
            if eid in id_map:
                props["id"] = id_map[eid]

        # ── 替换关系中的 start_node_id / end_node_id ──
        for rel in body.relationships:
            for key in ("start_node_id", "end_node_id", "subject", "object"):
                val = (rel.get(key) or "").strip()
                if val in id_map:
                    rel[key] = id_map[val]

        # ── 创建实体节点（按本体类型打标签）──
        for ent in body.entities:
            name = (ent.get("name") or "").strip()
            ont_type = (ent.get("ont_type") or "M_ENTITY").strip()
            props = ent.get("properties") or {}

            if not name:
                continue

            # 确保 name 在属性中
            if "name" not in props:
                props["name"] = name

            # ── 补全缺失字段 ──
            # 沿继承链（M_ROOT → … → ont_type）获取所有字段定义
            inherited = _get_inherited_fields(ont_type)
            if inherited:
                for code, finfo in inherited.items():
                    if code not in props or not props.get(code):
                        default_val = finfo.get("default") or ""
                        if default_val:
                            props[code] = default_val
                            filled_fields_count += 1
                        elif finfo.get("required", "0") == "1":
                            # 必填字段无默认值时，填入空字符串占位
                            props[code] = ""
                            filled_fields_count += 1

            # 清理 None 值
            clean_props = {k: v for k, v in props.items() if v is not None and v != ""}

            # Build dynamic label from ont_type (e.g. M_ENTITY -> Entity)
            type_to_label = {
                "M_ENTITY": "Entity",
                "M_BEHAVIOR": "Behavior",
                "M_RULE": "Rule",
                "M_SCENE": "Scene",
                "M_AGENT": "Agent",
                "M_EXCEPTION": "Exception",
                "M_QUALITY": "Quality",
                "M_EVENT": "Event",
                "M_TEMPLATE": "Template",
                "M_ROOT": "Entity",
                "M_BASE_ONTOLOGY": "Entity",
            }
            node_label = type_to_label.get(ont_type, "Entity")

            # 用 name 作为 merge key，创建带标签的节点
            await session.run(
                f"""
                MERGE (n:{node_label} {{name: $name}})
                SET n += $props
                SET n.ont_type = $ont_type
                """,
                name=name,
                props=clean_props,
                ont_type=ont_type,
            )

        # ── 创建关系 ──
        for rel in body.relationships:
            subj = (rel.get("start_node_id") or rel.get("subject") or "").strip()
            pred = (rel.get("type") or rel.get("predicate") or "").strip()
            obj = (rel.get("end_node_id") or rel.get("object") or "").strip()
            if not subj or not pred or not obj:
                continue

            safe_pred = pred.replace("`", "").replace(" ", "_")
            await session.run(
                f"""
                MATCH (a {{name: $subj}})
                MATCH (b {{name: $obj}})
                MERGE (a)-[r:RELATES {{type: $pred}}]->(b)
                SET r.predicate = $pred
                """,
                subj=subj, obj=obj, pred=pred,
            )
            edges_created += 1

        # ── 统计 ──
        node_result = await session.run("MATCH (n) RETURN count(n) AS cnt")
        node_rec = await node_result.single()
        nodes_created = node_rec["cnt"] if node_rec else 0

    # ── 绑定场景：实体名 ↔ 场景ID 写入 ontol_node_scene_relation ──
    scene_bind_count = 0
    if body.scene_ids:
        import uuid as _uuid
        for ent in body.entities:
            entity_name = (ent.get("name") or "").strip()
            if not entity_name:
                continue
            for sid in body.scene_ids:
                try:
                    await _execute_scene(
                        "INSERT INTO ontol_node_scene_relation (id, scene_id, scene_desc) VALUES (?,?,?)",
                        (_uuid.uuid4().hex[:16], sid, entity_name),
                    )
                    scene_bind_count += 1
                except Exception:
                    pass

    return {
        "filename": body.filename,
        "nodes_created": nodes_created,
        "edges_created": edges_created,
        "entity_count": len(body.entities),
        "filled_fields": filled_fields_count,
        "scene_bind_count": scene_bind_count,
    }


# 保留旧的 import-triples 兼容接口
class ImportTriplesRequest(BaseModel):
    filename: str = Field(..., description="来源文件名")
    triples: list[dict] = Field(..., description="三元组列表")


@router.post("/upload/import-triples")
async def import_triples_to_neo4j(body: ImportTriplesRequest):
    """兼容旧接口: 将三元组导入图数据库。新代码请使用 /upload/import-entities。"""
    return await import_entities_to_neo4j(
        ImportEntitiesRequest(
            filename=body.filename,
            entities=[
                {"name": t.get("subject", ""), "ont_type": "M_ENTITY", "properties": {"name": t.get("subject", "")}}
                for t in body.triples
            ]
            + [
                {"name": t.get("object", ""), "ont_type": "M_ENTITY", "properties": {"name": t.get("object", "")}}
                for t in body.triples
            ],
            relationships=body.triples,
        )
    )


# =========================================================================
# 场景管理 CRUD (ontol_model_scene)
# =========================================================================

class SceneCreate(BaseModel):
    id: str = Field(default="", max_length=32, description="场景ID，留空自动生成UUID")
    name: str = Field(..., max_length=100, description="场景名称")
    code: Optional[str] = Field(None, max_length=50, description="场景编码（唯一）")
    scene_desc: Optional[str] = Field(None, max_length=500, description="场景描述")
    parent_scene_id: Optional[str] = Field(None, max_length=32, description="父场景ID")
    create_id: Optional[str] = Field(None, max_length=32, description="创建人ID")

class SceneUpdate(BaseModel):
    name: Optional[str] = Field(None, max_length=100, description="场景名称")
    code: Optional[str] = Field(None, max_length=50, description="场景编码")
    scene_desc: Optional[str] = Field(None, max_length=500, description="场景描述")
    parent_scene_id: Optional[str] = Field(None, max_length=32, description="父场景ID")
    delete_flag: Optional[str] = Field(None, max_length=2, description="删除标志")


async def _query_scene(sql: str, params: tuple = ()) -> list[dict]:
    """执行场景表查询。"""
    import sqlite3
    db_path = _Path("infrastructure/db/ontol.db")
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    try:
        rows = conn.execute(sql, params).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


async def _execute_scene(sql: str, params: tuple = ()) -> None:
    """执行场景表写操作。"""
    import sqlite3
    db_path = _Path("infrastructure/db/ontol.db")
    conn = sqlite3.connect(str(db_path))
    try:
        conn.execute(sql, params)
        conn.commit()
    finally:
        conn.close()


@router.get("/scenes")
async def list_scenes(keyword: Optional[str] = None):
    """列出所有场景，支持关键词搜索。"""
    try:
        if keyword:
            rows = await _query_scene(
                "SELECT * FROM ontol_model_scene WHERE delete_flag='0' AND (name LIKE ? OR scene_desc LIKE ?) ORDER BY create_time DESC",
                (f"%{keyword}%", f"%{keyword}%"),
            )
        else:
            rows = await _query_scene(
                "SELECT * FROM ontol_model_scene WHERE delete_flag='0' ORDER BY create_time DESC"
            )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scene query failed: {e}")


@router.get("/scenes/{scene_id}")
async def get_scene(scene_id: str):
    """获取单个场景详情。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_model_scene WHERE id=? AND delete_flag='0'",
            (scene_id,),
        )
        if not rows:
            raise HTTPException(status_code=404, detail=f"Scene '{scene_id}' not found")
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scene query failed: {e}")


@router.post("/scenes", status_code=201)
async def create_scene(body: SceneCreate):
    """创建场景。id 留空时自动生成 UUID。"""
    import uuid as _uuid
    try:
        scene_id = body.id.strip() if body.id else _uuid.uuid4().hex[:16]
        existing = await _query_scene(
            "SELECT id FROM ontol_model_scene WHERE id=? AND delete_flag='0'",
            (scene_id,),
        )
        if existing:
            raise HTTPException(status_code=409, detail=f"Scene ID '{scene_id}' already exists")

        # 检查 code 唯一性
        if body.code:
            dup = await _query_scene(
                "SELECT id FROM ontol_model_scene WHERE code=? AND delete_flag='0'",
                (body.code,),
            )
            if dup:
                raise HTTPException(status_code=409, detail=f"场景编码 '{body.code}' 已被使用")

        await _execute_scene(
            "INSERT INTO ontol_model_scene (id, name, code, scene_desc, parent_scene_id, create_id) VALUES (?,?,?,?,?,?)",
            (scene_id, body.name, body.code or None, body.scene_desc or "", body.parent_scene_id or None, body.create_id or ""),
        )
        return await get_scene(scene_id)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scene create failed: {e}")


@router.put("/scenes/{scene_id}")
async def update_scene(scene_id: str, body: SceneUpdate):
    """更新场景。系统预设场景（scene_is_system='1'）不可修改。"""
    try:
        rows = await _query_scene(
            "SELECT id, scene_is_system FROM ontol_model_scene WHERE id=? AND delete_flag='0'",
            (scene_id,),
        )
        if not rows:
            raise HTTPException(status_code=404, detail=f"Scene '{scene_id}' not found")
        if rows[0].get("scene_is_system") == "1":
            raise HTTPException(status_code=403, detail="系统预设场景不可修改")

        # 构建动态 UPDATE
        sets, params = [], []
        if body.name is not None:
            sets.append("name=?")
            params.append(body.name)
        if body.scene_desc is not None:
            sets.append("scene_desc=?")
            params.append(body.scene_desc)
        if body.code is not None:
            if body.code:
                dup = await _query_scene(
                    "SELECT id FROM ontol_model_scene WHERE code=? AND delete_flag='0' AND id!=?",
                    (body.code, scene_id),
                )
                if dup:
                    raise HTTPException(status_code=409, detail=f"场景编码 '{body.code}' 已被使用")
            sets.append("code=?")
            params.append(body.code)
        if body.parent_scene_id is not None:
            sets.append("parent_scene_id=?")
            params.append(body.parent_scene_id if body.parent_scene_id else None)
        if body.delete_flag is not None:
            sets.append("delete_flag=?")
            params.append(body.delete_flag)
        if not sets:
            return await get_scene(scene_id)

        params.append(scene_id)
        await _execute_scene(
            f"UPDATE ontol_model_scene SET {', '.join(sets)} WHERE id=?",
            tuple(params),
        )
        return await get_scene(scene_id)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scene update failed: {e}")


@router.delete("/scenes/{scene_id}")
async def delete_scene(scene_id: str, soft: bool = True):
    """删除场景（默认软删除）。系统预设场景（scene_is_system='1'）不可删除。"""
    try:
        # 检查是否为系统预设
        rows = await _query_scene(
            "SELECT scene_is_system FROM ontol_model_scene WHERE id=? AND delete_flag='0'",
            (scene_id,),
        )
        if rows and rows[0].get("scene_is_system") == "1":
            raise HTTPException(status_code=403, detail="系统预设场景不可删除")
        if soft:
            await _execute_scene(
                "UPDATE ontol_model_scene SET delete_flag='1' WHERE id=?",
                (scene_id,),
            )
        else:
            await _execute_scene(
                "DELETE FROM ontol_model_scene WHERE id=?",
                (scene_id,),
            )
        return {"deleted": True, "scene_id": scene_id, "soft": soft}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scene delete failed: {e}")


# =========================================================================
# 对话-场景关系 CRUD (ontol_char_scene_relation)
# =========================================================================

class ChatSceneBind(BaseModel):
    chat_id: str = Field(..., description="对话UUID")
    scene_ids: list[str] = Field(..., description="场景ID列表")


@router.post("/chat-scenes/bind")
async def bind_chat_scenes(body: ChatSceneBind):
    """将对话绑定到多个场景（先删旧绑定，再批量插入）。"""
    try:
        # 先清理旧绑定
        await _execute_scene(
            "UPDATE ontol_char_scene_relation SET delete_flag='1' WHERE chat_id=?",
            (body.chat_id,),
        )
        # 批量插入新绑定
        import uuid as _uuid
        for sid in body.scene_ids:
            rid = _uuid.uuid4().hex[:16]
            await _execute_scene(
                "INSERT INTO ontol_char_scene_relation (id, scene_id) VALUES (?,?,?)",
                (rid, body.sid),
            )
        return await get_chat_scenes(body.chat_id)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Bind chat scenes failed: {e}")


@router.get("/chat-scenes/{chat_id}")
async def get_chat_scenes(chat_id: str):
    """获取对话绑定的场景列表（带场景名称）。"""
    try:
        rows = await _query_scene(
            """SELECT r.id, r.r.scene_id, s.name, s.scene_desc
               FROM ontol_char_scene_relation r
               LEFT JOIN ontol_model_scene s ON r.scene_id = s.id
               WHERE r.chat_id=? AND r.delete_flag='0'
               ORDER BY r.create_time""",
            (chat_id,),
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query chat scenes failed: {e}")


@router.delete("/chat-scenes/{relation_id}")
async def unbind_chat_scene(relation_id: str):
    """删除对话-场景绑定（软删除）。"""
    try:
        await _execute_scene(
            "UPDATE ontol_char_scene_relation SET delete_flag='1' WHERE id=?",
            (relation_id,),
        )
        return {"deleted": True, "relation_id": relation_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unbind chat scene failed: {e}")


# =========================================================================
# 提示词管理 CRUD (ontol_scene_prompt)
# =========================================================================

class ScenePromptCreate(BaseModel):
    id: str = Field(..., max_length=32, description="提示词ID（主键）")
    scene_id: str = Field(..., max_length=32, description="所属场景ID")
    name: str = Field(..., max_length=100, description="提示词名称")
    prompt_content: str = Field(default="", description="提示词文本内容")
    prompt_desc: Optional[str] = Field(None, max_length=500, description="提示词描述")
    prompt_description: Optional[str] = Field(None, max_length=500, description="提示词调用时机说明")
    create_id: Optional[str] = Field(None, max_length=32, description="创建人ID")

class ScenePromptUpdate(BaseModel):
    name: Optional[str] = Field(None, max_length=100)
    prompt_content: Optional[str] = None
    prompt_desc: Optional[str] = Field(None, max_length=500)
    prompt_description: Optional[str] = Field(None, max_length=500)
    delete_flag: Optional[str] = Field(None, max_length=2)


@router.get("/scenes/{scene_id}/prompts")
async def list_scene_prompts(scene_id: str):
    """列出场景下的所有提示词。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_scene_prompt WHERE scene_id=? AND delete_flag='0' ORDER BY create_time DESC",
            (scene_id,),
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query prompts failed: {e}")


@router.get("/prompts/{prompt_id}")
async def get_prompt(prompt_id: str):
    """获取单个提示词详情。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_scene_prompt WHERE id=? AND delete_flag='0'",
            (prompt_id,),
        )
        if not rows:
            raise HTTPException(status_code=404, detail=f"Prompt {prompt_id} not found")
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query prompt failed: {e}")


@router.post("/scenes/{scene_id}/prompts", status_code=201)
async def create_prompt(scene_id: str, body: ScenePromptCreate):
    """创建提示词。后端生成 UUID 主键。"""
    import uuid as _uuid
    try:
        prompt_id = body.id if body.id else _uuid.uuid4().hex[:16]
        existing = await _query_scene("SELECT id FROM ontol_scene_prompt WHERE id=?", (prompt_id,))
        if existing:
            raise HTTPException(status_code=409, detail=f"Prompt ID '{prompt_id}' already exists")
        await _execute_scene(
            "INSERT INTO ontol_scene_prompt (id, scene_id, name, prompt_content, prompt_desc, prompt_description, create_id) VALUES (?,?,?,?,?,?,?)",
            (prompt_id, scene_id, body.name, body.prompt_content or "", body.prompt_desc or "", body.prompt_description or "", body.create_id or ""),
        )
        rows = await _query_scene("SELECT * FROM ontol_scene_prompt WHERE id=?", (prompt_id,))
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Create prompt failed: {e}")


@router.put("/prompts/{prompt_id}")
async def update_prompt(prompt_id: str, body: ScenePromptUpdate):
    """更新提示词。"""
    try:
        existing = await _query_scene("SELECT id FROM ontol_scene_prompt WHERE id=? AND delete_flag='0'", (prompt_id,))
        if not existing:
            raise HTTPException(status_code=404, detail=f"Prompt {prompt_id} not found")
        data = body.model_dump(exclude_none=True)
        if data:
            set_clause = ", ".join(f"{k}=?" for k in data)
            values = list(data.values()) + [prompt_id]
            await _execute_scene(f"UPDATE ontol_scene_prompt SET {set_clause} WHERE id=?", tuple(values))
        rows = await _query_scene("SELECT * FROM ontol_scene_prompt WHERE id=?", (prompt_id,))
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Update prompt failed: {e}")


@router.delete("/prompts/{prompt_id}")
async def delete_prompt(prompt_id: str, soft: bool = True):
    """删除提示词（默认软删除）。"""
    try:
        if soft:
            await _execute_scene("UPDATE ontol_scene_prompt SET delete_flag='1' WHERE id=?", (prompt_id,))
        else:
            await _execute_scene("DELETE FROM ontol_scene_prompt WHERE id=?", (prompt_id,))
        return {"deleted": True, "prompt_id": prompt_id, "soft": soft}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete prompt failed: {e}")


# =========================================================================
# 场景词典 CRUD (ontol_scene_dictionary)
# =========================================================================

class SceneDictCreate(BaseModel):
    id: str = Field(default="", max_length=32, description="词典ID，留空自动生成UUID")
    scene_id: str = Field(..., max_length=32, description="所属场景ID")
    name: str = Field(..., max_length=200, description="词典名称")
    code: str = Field(..., max_length=100, description="词典编码（必填唯一）")
    dictionary_type_id: Optional[str] = Field(None, max_length=32, description="词条分类ID")
    dictionary_content: Optional[str] = Field(None, description="词典内容")

class SceneDictUpdate(BaseModel):
    name: Optional[str] = Field(None, max_length=200)
    code: Optional[str] = Field(None, max_length=100)
    dictionary_type_id: Optional[str] = Field(None, max_length=32)
    dictionary_content: Optional[str] = None
    delete_flag: Optional[str] = Field(None, max_length=2)


@router.get("/scenes/{scene_id}/dictionaries")
async def list_scene_dicts(scene_id: str):
    """列出场景下的所有词典，连带词条分类名。"""
    try:
        rows = await _query_scene(
            """SELECT d.*, dt.name
               FROM ontol_scene_dictionary d
               LEFT JOIN ontol_dictionary_type dt ON d.dictionary_type_id = dt.id
               WHERE d.scene_id=? AND d.delete_flag='0'
               ORDER BY d.create_time DESC""",
            (scene_id,),
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query dictionaries failed: {e}")


@router.get("/dictionaries")
async def list_all_dicts(dictionary_type_id: str = ""):
    """列出所有词典（可选的按词条分类筛选），带场景名和分类名。"""
    try:
        if dictionary_type_id:
            rows = await _query_scene(
                """SELECT d.*, s.name, dt.name
                   FROM ontol_scene_dictionary d
                   LEFT JOIN ontol_model_scene s ON d.scene_id = s.id
                   LEFT JOIN ontol_dictionary_type dt ON d.dictionary_type_id = dt.id
                   WHERE d.delete_flag='0' AND d.dictionary_type_id=?
                   ORDER BY d.create_time DESC""",
                (dictionary_type_id,),
            )
        else:
            rows = await _query_scene(
                """SELECT d.*, s.name, dt.name
                   FROM ontol_scene_dictionary d
                   LEFT JOIN ontol_model_scene s ON d.scene_id = s.id
                   LEFT JOIN ontol_dictionary_type dt ON d.dictionary_type_id = dt.id
                   WHERE d.delete_flag='0'
                   ORDER BY d.create_time DESC"""
            )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query dictionaries failed: {e}")


@router.post("/dictionaries", status_code=201)
async def create_dict_direct(body: SceneDictCreate):
    """创建词典（直发接口，scene_id 从 body 传入）。"""
    import uuid as _uuid
    try:
        dict_id = body.id.strip() if body.id else _uuid.uuid4().hex[:16]
        existing = await _query_scene("SELECT id FROM ontol_scene_dictionary WHERE id=?", (dict_id,))
        if existing:
            raise HTTPException(status_code=409, detail=f"Dictionary ID '{dict_id}' already exists")
        await _execute_scene(
            "INSERT INTO ontol_scene_dictionary (id, scene_id, name, code, dictionary_type_id, dictionary_content) VALUES (?,?,?,?,?,?)",
            (dict_id, body.scene_id, body.name, body.code, body.dictionary_type_id or None, body.dictionary_content or ""),
        )
        rows = await _query_scene("SELECT * FROM ontol_scene_dictionary WHERE id=?", (dict_id,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Create dictionary failed: {e}")


@router.get("/dictionaries/{dict_id}")
async def get_dict(dict_id: str):
    """获取单个词典详情。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_scene_dictionary WHERE id=? AND delete_flag='0'",
            (dict_id,),
        )
        if not rows:
            raise HTTPException(status_code=404, detail=f"Dictionary {dict_id} not found")
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query dictionary failed: {e}")


@router.post("/scenes/{scene_id}/dictionaries", status_code=201)
async def create_dict(scene_id: str, body: SceneDictCreate):
    """创建场景词典，id 留空自动生成 UUID。"""
    import uuid as _uuid
    try:
        dict_id = body.id.strip() if body.id else _uuid.uuid4().hex[:16]
        existing = await _query_scene("SELECT id FROM ontol_scene_dictionary WHERE id=?", (dict_id,))
        if existing:
            raise HTTPException(status_code=409, detail=f"Dictionary ID '{dict_id}' already exists")
        await _execute_scene(
            "INSERT INTO ontol_scene_dictionary (id, scene_id, name, code, dictionary_type_id, dictionary_content) VALUES (?,?,?,?,?,?)",
            (dict_id, scene_id, body.name, body.code, body.dictionary_type_id or None, body.dictionary_content or ""),
        )
        rows = await _query_scene("SELECT * FROM ontol_scene_dictionary WHERE id=?", (dict_id,))
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Create dictionary failed: {e}")


@router.put("/dictionaries/{dict_id}")
async def update_dict(dict_id: str, body: SceneDictUpdate):
    """更新场景词典。"""
    try:
        existing = await _query_scene("SELECT id FROM ontol_scene_dictionary WHERE id=? AND delete_flag='0'", (dict_id,))
        if not existing:
            raise HTTPException(status_code=404, detail=f"Dictionary {dict_id} not found")
        data = body.model_dump(exclude_none=True)
        if data:
            set_clause = ", ".join(f"{k}=?" for k in data)
            values = list(data.values()) + [dict_id]
            await _execute_scene(f"UPDATE ontol_scene_dictionary SET {set_clause} WHERE id=?", tuple(values))
        rows = await _query_scene("SELECT * FROM ontol_scene_dictionary WHERE id=?", (dict_id,))
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Update dictionary failed: {e}")


@router.delete("/dictionaries/{dict_id}")
async def delete_dict(dict_id: str, soft: bool = True):
    """删除场景词典（默认软删除）。"""
    try:
        if soft:
            await _execute_scene("UPDATE ontol_scene_dictionary SET delete_flag='1' WHERE id=?", (dict_id,))
        else:
            await _execute_scene("DELETE FROM ontol_scene_dictionary WHERE id=?", (dict_id,))
        return {"deleted": True, "dict_id": dict_id, "soft": soft}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete dictionary failed: {e}")


# =========================================================================
# 词条分类 CRUD (ontol_dictionary_type)
# =========================================================================

class DictTypeCreate(BaseModel):
    id: str = Field(default="", max_length=32, description="UUID")
    name: str = Field(..., max_length=200, description="分类名称")
    dictionary_description: Optional[str] = Field(None, max_length=500)
    is_system: str = Field(default="0", max_length=1, description="是否系统预设")

class DictTypeUpdate(BaseModel):
    name: Optional[str] = Field(None, max_length=200)
    dictionary_description: Optional[str] = Field(None, max_length=500)
    delete_flag: Optional[str] = Field(None, max_length=2)


@router.get("/dictionary-types")
async def list_dict_types():
    """列出所有词条分类。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_dictionary_type WHERE delete_flag='0' ORDER BY create_time DESC"
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query dictionary types failed: {e}")


@router.get("/dictionary-types/{type_id}")
async def get_dict_type(type_id: str):
    """获取单个词条分类。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_dictionary_type WHERE id=? AND delete_flag='0'", (type_id,),
        )
        if not rows: raise HTTPException(status_code=404, detail="Not found")
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Query failed: {e}")


@router.post("/dictionary-types", status_code=201)
async def create_dict_type(body: DictTypeCreate):
    """创建词条分类。"""
    import uuid as _uuid
    try:
        tid = body.id.strip() if body.id else _uuid.uuid4().hex[:16]
        await _execute_scene(
            "INSERT INTO ontol_dictionary_type (id, name, dictionary_description, is_system) VALUES (?,?,?,?)",
            (tid, body.name, body.dictionary_description or "", body.is_system),
        )
        rows = await _query_scene("SELECT * FROM ontol_dictionary_type WHERE id=?", (tid,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Create failed: {e}")


@router.put("/dictionary-types/{type_id}")
async def update_dict_type(type_id: str, body: DictTypeUpdate):
    """更新词条分类。系统预设不可修改。"""
    try:
        existing = await _query_scene("SELECT * FROM ontol_dictionary_type WHERE id=? AND delete_flag='0'", (type_id,))
        if not existing: raise HTTPException(status_code=404, detail="Not found")
        if existing[0].get("is_system") == "1":
            raise HTTPException(status_code=403, detail="系统预设分类不可修改")
        data = body.model_dump(exclude_none=True)
        if data:
            sets = ", ".join(f"{k}=?" for k in data)
            await _execute_scene(f"UPDATE ontol_dictionary_type SET {sets} WHERE id=?", tuple(data.values()) + (type_id,))
        rows = await _query_scene("SELECT * FROM ontol_dictionary_type WHERE id=?", (type_id,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Update failed: {e}")


@router.delete("/dictionary-types/{type_id}")
async def delete_dict_type(type_id: str, soft: bool = True):
    """删除词条分类。系统预设不可删除。"""
    try:
        existing = await _query_scene("SELECT is_system FROM ontol_dictionary_type WHERE id=? AND delete_flag='0'", (type_id,))
        if not existing: raise HTTPException(status_code=404, detail="Not found")
        if existing[0].get("is_system") == "1":
            raise HTTPException(status_code=403, detail="系统预设分类不可删除")
        if soft:
            await _execute_scene("UPDATE ontol_dictionary_type SET delete_flag='1' WHERE id=?", (type_id,))
        else:
            await _execute_scene("DELETE FROM ontol_dictionary_type WHERE id=?", (type_id,))
        return {"deleted": True, "type_id": type_id, "soft": soft}
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Delete failed: {e}")


# =========================================================================
# LLM 类型配置 CRUD (ontol_llm_type_config)
# =========================================================================

class LLMTypeConfigCreate(BaseModel):
    id: str = Field(default="", max_length=32, description="UUID")
    name: str = Field(..., max_length=200, description="类型名称")
    llm_description: Optional[str] = Field(None, max_length=500)
    is_system: str = Field(default="0", max_length=1, description="是否系统预设 0自定义/1系统")

class LLMTypeConfigUpdate(BaseModel):
    name: Optional[str] = Field(None, max_length=200)
    llm_description: Optional[str] = Field(None, max_length=500)
    delete_flag: Optional[str] = Field(None, max_length=2)


@router.get("/llm-type-configs")
async def list_llm_type_configs():
    """列出所有 LLM 类型配置。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_llm_type_config WHERE delete_flag='0' ORDER BY create_time DESC"
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query failed: {e}")


@router.get("/llm-type-configs/{config_id}")
async def get_llm_type_config(config_id: str):
    """获取单个类型配置。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_llm_type_config WHERE id=? AND delete_flag='0'", (config_id,),
        )
        if not rows: raise HTTPException(status_code=404, detail=f"Not found")
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Query failed: {e}")


@router.post("/llm-type-configs", status_code=201)
async def create_llm_type_config(body: LLMTypeConfigCreate):
    """创建 LLM 类型配置。"""
    import uuid as _uuid
    try:
        cid = body.id.strip() if body.id else _uuid.uuid4().hex[:16]
        await _execute_scene(
            "INSERT INTO ontol_llm_type_config (id, name, llm_description, is_system) VALUES (?,?,?,?)",
            (cid, body.name, body.llm_description or "", body.is_system),
        )
        rows = await _query_scene("SELECT * FROM ontol_llm_type_config WHERE id=?", (cid,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Create failed: {e}")


@router.put("/llm-type-configs/{config_id}")
async def update_llm_type_config(config_id: str, body: LLMTypeConfigUpdate):
    """更新类型配置。系统预设（is_system='1'）不可修改。"""
    try:
        existing = await _query_scene("SELECT * FROM ontol_llm_type_config WHERE id=? AND delete_flag='0'", (config_id,))
        if not existing: raise HTTPException(status_code=404, detail=f"Not found")
        if existing[0].get("is_system") == "1":
            raise HTTPException(status_code=403, detail="系统预设分类不可修改")
        data = body.model_dump(exclude_none=True)
        if data:
            sets = ", ".join(f"{k}=?" for k in data)
            await _execute_scene(f"UPDATE ontol_llm_type_config SET {sets} WHERE id=?", tuple(data.values()) + (config_id,))
        rows = await _query_scene("SELECT * FROM ontol_llm_type_config WHERE id=?", (config_id,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Update failed: {e}")


@router.delete("/llm-type-configs/{config_id}")
async def delete_llm_type_config(config_id: str, soft: bool = True):
    """删除类型配置。系统预设（is_system='1'）不可删除。"""
    try:
        existing = await _query_scene("SELECT is_system FROM ontol_llm_type_config WHERE id=? AND delete_flag='0'", (config_id,))
        if not existing: raise HTTPException(status_code=404, detail=f"Not found")
        if existing[0].get("is_system") == "1":
            raise HTTPException(status_code=403, detail="系统预设分类不可删除")
        if soft:
            await _execute_scene("UPDATE ontol_llm_type_config SET delete_flag='1' WHERE id=?", (config_id,))
        else:
            await _execute_scene("DELETE FROM ontol_llm_type_config WHERE id=?", (config_id,))
        return {"deleted": True, "config_id": config_id, "soft": soft}
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Delete failed: {e}")


# =========================================================================
# LLM 模型配置 CRUD (ontol_llm_config)
# =========================================================================

class LLMConfigCreate(BaseModel):
    id: str = Field(default="", max_length=32, description="UUID")
    llm_type_config_id: Optional[str] = Field(None, max_length=32, description="所属类型配置ID")
    name: str = Field(..., max_length=200, description="显示名")
    llm_model: Optional[str] = Field(None, max_length=200, description="API模型名（如 deepseek-v4-pro）")
    llm_url: Optional[str] = Field(None, max_length=500, description="调用地址")
    llm_key: Optional[str] = Field(None, max_length=500, description="调用Key")
    llm_description: Optional[str] = Field(None, max_length=500, description="描述说明")

class LLMConfigUpdate(BaseModel):
    llm_type_config_id: Optional[str] = Field(None, max_length=32)
    name: Optional[str] = Field(None, max_length=200)
    llm_model: Optional[str] = Field(None, max_length=200)
    llm_url: Optional[str] = Field(None, max_length=500)
    llm_key: Optional[str] = Field(None, max_length=500)
    llm_description: Optional[str] = Field(None, max_length=500)
    delete_flag: Optional[str] = Field(None, max_length=2)


@router.get("/llm-configs")
async def list_llm_configs(type_config_id: Optional[str] = None):
    """列出 LLM 配置，可按类型配置ID筛选。"""
    try:
        if type_config_id:
            rows = await _query_scene(
                "SELECT * FROM ontol_llm_config WHERE delete_flag='0' AND llm_type_config_id=? ORDER BY create_time DESC",
                (type_config_id,),
            )
        else:
            rows = await _query_scene(
                "SELECT * FROM ontol_llm_config WHERE delete_flag='0' ORDER BY create_time DESC"
            )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query llm configs failed: {e}")


@router.get("/llm-configs/{config_id}")
async def get_llm_config(config_id: str):
    """获取单个 LLM 配置详情。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_llm_config WHERE id=? AND delete_flag='0'",
            (config_id,),
        )
        if not rows:
            raise HTTPException(status_code=404, detail=f"Config {config_id} not found")
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query config failed: {e}")


@router.post("/llm-configs", status_code=201)
async def create_llm_config(body: LLMConfigCreate):
    """创建 LLM 配置，id 留空自动 UUID。"""
    import uuid as _uuid
    try:
        cfg_id = body.id.strip() if body.id else _uuid.uuid4().hex[:16]
        existing = await _query_scene("SELECT id FROM ontol_llm_config WHERE id=?", (cfg_id,))
        if existing:
            raise HTTPException(status_code=409, detail=f"Config ID '{cfg_id}' already exists")
        await _execute_scene(
            "INSERT INTO ontol_llm_config (id, llm_type_config_id, name, llm_model, llm_url, llm_key, llm_description) VALUES (?,?,?,?,?,?,?)",
            (cfg_id, body.llm_type_config_id or None, body.name, body.llm_model or None, body.llm_url or "", body.llm_key or "", body.llm_description or ""),
        )
        rows = await _query_scene("SELECT * FROM ontol_llm_config WHERE id=?", (cfg_id,))
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Create config failed: {e}")


@router.put("/llm-configs/{config_id}")
async def update_llm_config(config_id: str, body: LLMConfigUpdate):
    """更新 LLM 配置。"""
    try:
        existing = await _query_scene("SELECT id FROM ontol_llm_config WHERE id=? AND delete_flag='0'", (config_id,))
        if not existing:
            raise HTTPException(status_code=404, detail=f"Config {config_id} not found")
        data = body.model_dump(exclude_none=True)
        if data:
            set_clause = ", ".join(f"{k}=?" for k in data)
            values = list(data.values()) + [config_id]
            await _execute_scene(f"UPDATE ontol_llm_config SET {set_clause} WHERE id=?", tuple(values))
        rows = await _query_scene("SELECT * FROM ontol_llm_config WHERE id=?", (config_id,))
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Update config failed: {e}")


@router.delete("/llm-configs/{config_id}")
async def delete_llm_config(config_id: str, soft: bool = True):
    """删除 LLM 配置（默认软删除）。"""
    try:
        if soft:
            await _execute_scene("UPDATE ontol_llm_config SET delete_flag='1' WHERE id=?", (config_id,))
        else:
            await _execute_scene("DELETE FROM ontol_llm_config WHERE id=?", (config_id,))
        return {"deleted": True, "config_id": config_id, "soft": soft}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete config failed: {e}")


# =========================================================================
# 动态函数类型 CRUD (ontol_function_type)
# =========================================================================

class FunctionTypeCreate(BaseModel):
    id: str = Field(default="", max_length=32)
    name: str = Field(..., max_length=200)
    function_description: Optional[str] = Field(None, max_length=500)
    is_system: Optional[str] = Field("0", max_length=2)

class FunctionTypeUpdate(BaseModel):
    name: Optional[str] = Field(None, max_length=200)
    function_description: Optional[str] = Field(None, max_length=500)
    is_system: Optional[str] = Field(None, max_length=2)
    delete_flag: Optional[str] = Field(None, max_length=2)


@router.get("/function-types")
async def list_function_types():
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_function_type WHERE delete_flag='0' ORDER BY create_time DESC"
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query failed: {e}")

@router.get("/function-types/{type_id}")
async def get_function_type(type_id: str):
    try:
        rows = await _query_scene("SELECT * FROM ontol_function_type WHERE id=? AND delete_flag='0'", (type_id,))
        if not rows: raise HTTPException(status_code=404, detail=f"Type {type_id} not found")
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Query failed: {e}")

@router.post("/function-types", status_code=201)
async def create_function_type(body: FunctionTypeCreate):
    import uuid as _uuid
    try:
        tid = body.id.strip() if body.id else _uuid.uuid4().hex[:16]
        await _execute_scene(
            "INSERT INTO ontol_function_type (id, name, function_description, is_system) VALUES (?,?,?,?)",
            (tid, body.name, body.function_description or None, body.is_system or "0"),
        )
        rows = await _query_scene("SELECT * FROM ontol_function_type WHERE id=?", (tid,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Create failed: {e}")

@router.put("/function-types/{type_id}")
async def update_function_type(type_id: str, body: FunctionTypeUpdate):
    try:
        existing = await _query_scene("SELECT id FROM ontol_function_type WHERE id=? AND delete_flag='0'", (type_id,))
        if not existing: raise HTTPException(status_code=404, detail=f"Type {type_id} not found")
        data = body.model_dump(exclude_none=True)
        if data:
            sets = ", ".join(f"{k}=?" for k in data)
            await _execute_scene(f"UPDATE ontol_function_type SET {sets} WHERE id=?", tuple(data.values()) + (type_id,))
        rows = await _query_scene("SELECT * FROM ontol_function_type WHERE id=?", (type_id,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Update failed: {e}")

@router.delete("/function-types/{type_id}")
async def delete_function_type(type_id: str, soft: bool = True):
    try:
        if soft:
            await _execute_scene("UPDATE ontol_function_type SET delete_flag='1' WHERE id=?", (type_id,))
        else:
            await _execute_scene("DELETE FROM ontol_function_type WHERE id=?", (type_id,))
        return {"deleted": True, "type_id": type_id, "soft": soft}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete failed: {e}")


# =========================================================================
# 动态函数 CRUD (ontol_function)
# =========================================================================

class FunctionCreate(BaseModel):
    id: str = Field(default="", max_length=32)
    function_type_id: Optional[str] = Field(None, max_length=32)
    code: str = Field(..., max_length=100, description="函数唯一标识")
    name: str = Field(..., max_length=100, description="函数名称简写")
    function_classpath: Optional[str] = Field(None, max_length=255)
    function_method: Optional[str] = Field(None, max_length=100)
    function_type: Optional[str] = Field("PYTHON", max_length=20)
    function_timeout_ms: Optional[int] = 30000
    function_max_retry: Optional[int] = 0
    status: Optional[int] = 1
    description: Optional[str] = None

class FunctionUpdate(BaseModel):
    function_type_id: Optional[str] = Field(None, max_length=32)
    code: Optional[str] = Field(None, max_length=100)
    name: Optional[str] = Field(None, max_length=100)
    function_classpath: Optional[str] = Field(None, max_length=255)
    function_method: Optional[str] = Field(None, max_length=100)
    function_type: Optional[str] = Field(None, max_length=20)
    function_timeout_ms: Optional[int] = None
    function_max_retry: Optional[int] = None
    status: Optional[int] = None
    description: Optional[str] = None
    delete_flag: Optional[str] = Field(None, max_length=2)


@router.get("/functions")
async def list_functions():
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_function WHERE delete_flag='0' ORDER BY create_time DESC"
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query failed: {e}")

@router.get("/functions/{func_id}")
async def get_function(func_id: str):
    try:
        rows = await _query_scene("SELECT * FROM ontol_function WHERE id=? AND delete_flag='0'", (func_id,))
        if not rows: raise HTTPException(status_code=404, detail=f"Function {func_id} not found")
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Query failed: {e}")

@router.post("/functions", status_code=201)
async def create_function(body: FunctionCreate):
    import uuid as _uuid
    try:
        fid = body.id.strip() if body.id else _uuid.uuid4().hex[:16]
        await _execute_scene(
            "INSERT INTO ontol_function (id, function_type_id, code, name, function_classpath, function_method, function_type, function_timeout_ms, function_max_retry, status, description) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (fid, body.function_type_id or None, body.code, body.name, body.function_classpath or None, body.function_method or None, body.function_type or "PYTHON", body.function_timeout_ms or 30000, body.function_max_retry or 0, body.status or 1, body.description or None),
        )
        rows = await _query_scene("SELECT * FROM ontol_function WHERE id=?", (fid,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Create failed: {e}")

@router.put("/functions/{func_id}")
async def update_function(func_id: str, body: FunctionUpdate):
    try:
        existing = await _query_scene("SELECT id FROM ontol_function WHERE id=? AND delete_flag='0'", (func_id,))
        if not existing: raise HTTPException(status_code=404, detail=f"Function {func_id} not found")
        data = body.model_dump(exclude_none=True)
        if data:
            sets = ", ".join(f"{k}=?" for k in data)
            await _execute_scene(f"UPDATE ontol_function SET {sets} WHERE id=?", tuple(data.values()) + (func_id,))
        rows = await _query_scene("SELECT * FROM ontol_function WHERE id=?", (func_id,))
        return rows[0]
    except HTTPException: raise
    except Exception as e: raise HTTPException(status_code=500, detail=f"Update failed: {e}")

@router.delete("/functions/{func_id}")
async def delete_function(func_id: str, soft: bool = True):
    try:
        if soft:
            await _execute_scene("UPDATE ontol_function SET delete_flag='1' WHERE id=?", (func_id,))
        else:
            await _execute_scene("DELETE FROM ontol_function WHERE id=?", (func_id,))
        return {"deleted": True, "func_id": func_id, "soft": soft}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete failed: {e}")


# =========================================================================
# 推演版本管理 CRUD (ontol_cope_version)
# =========================================================================

class CopeVersionCreate(BaseModel):
    name: str = Field(default="", description="副本名称")
    cope_version_status: str = Field(default="00", description="状态: 00待处理/01推理中/02推理完成/03已删除")
    init_note_id: str = Field(default="", description="初始节点ID")
    init_note_name: str = Field(default="", description="初始节点名称")
    description: str = Field(default="", description="描述")
    confidence: float = Field(default=0.8, ge=0.01, le=1.0, description="置信度 (0.01~1.00)")

class CopeVersionUpdate(BaseModel):
    name: Optional[str] = Field(None, description="副本名称")
    cope_version_status: Optional[str] = Field(None, description="状态")
    init_note_id: Optional[str] = Field(None, description="初始节点ID")
    init_note_name: Optional[str] = Field(None, description="初始节点名称")
    description: Optional[str] = Field(None, description="描述")
    confidence: Optional[float] = Field(None, ge=0.01, le=1.0, description="置信度 (0.01~1.00)")


@router.get("/cope-versions")
async def list_cope_versions():
    """查询 ontol_cope_version 所有有效记录。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_cope_version WHERE delete_flag='0' ORDER BY create_time DESC"
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query cope versions failed: {e}")


@router.post("/cope-versions")
async def create_cope_version(body: CopeVersionCreate):
    """新增副本版本记录。"""
    import uuid as _uuid
    from datetime import datetime as _dt
    rid = _uuid.uuid4().hex[:16]
    now = _dt.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    try:
        await _execute_scene(
            "INSERT INTO ontol_cope_version "
            "(id, name, code, cope_version_status, description, init_note_id, init_note_name, confidence, create_time) "
            "VALUES (?,?,?,?,?,?,?,?,?)",
            (rid, body.name, rid, body.cope_version_status,
             body.description, body.init_note_id, body.init_note_name, body.confidence, now),
        )
        rows = await _query_scene("SELECT * FROM ontol_cope_version WHERE id=?", (rid,))
        return rows[0] if rows else {"id": rid}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Create cope version failed: {e}")


@router.put("/cope-versions/{cope_id}")
async def update_cope_version(cope_id: str, body: CopeVersionUpdate):
    """更新副本版本记录。"""
    from datetime import datetime as _dt
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")
    updates["update_time"] = _dt.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    try:
        sets = ", ".join(f"{k}=?" for k in updates)
        params = tuple(updates.values()) + (cope_id,)
        await _execute_scene(
            f"UPDATE ontol_cope_version SET {sets} WHERE id=?",
            params,
        )
        rows = await _query_scene("SELECT * FROM ontol_cope_version WHERE id=?", (cope_id,))
        return rows[0] if rows else {"id": cope_id, "updated": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Update cope version failed: {e}")


@router.delete("/cope-versions/{cope_id}")
async def delete_cope_version(cope_id: str):
    """软删除副本版本记录。"""
    from datetime import datetime as _dt
    try:
        now = _dt.utcnow().strftime("%Y-%m-%d %H:%M:%S")
        await _execute_scene(
            "UPDATE ontol_cope_version SET delete_flag='1', update_time=? WHERE id=?",
            (now, cope_id),
        )
        return {"deleted": True, "cope_id": cope_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete cope version failed: {e}")


@router.delete("/cope-versions/{cope_id}/nodes")
async def delete_cope_version_nodes(cope_id: str):
    """批量删除图数据库中 copy_version 匹配该记录 ID 的节点。"""
    from infrastructure.db.neo4j import get_driver

    try:
        driver = await get_driver()
        async with driver.session() as session:
            result = await session.run(
                "MATCH (n) WHERE n.copy_version = $cope_id DETACH DELETE n RETURN count(n) AS deleted",
                cope_id=cope_id,
            )
            record = await result.single()
            count = record["deleted"] if record else 0
        return {"deleted": True, "cope_id": cope_id, "node_count": count}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete cope version nodes failed: {e}")


@router.get("/cope-versions/{cope_id}")
async def get_cope_version(cope_id: str):
    """查询单条副本版本记录。"""
    try:
        rows = await _query_scene(
            "SELECT * FROM ontol_cope_version WHERE id=? AND delete_flag='0'",
            (cope_id,),
        )
        if not rows:
            raise HTTPException(status_code=404, detail="Cope version not found")
        return rows[0]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query cope version failed: {e}")


@router.get("/cope-versions/{cope_id}/graph")
async def get_cope_graph(cope_id: str):
    """
    获取推演副本的图数据（节点+关系）。
    - status=00: 返回没有 copy_version 属性的节点
    - 其他状态: 返回 copy_version=cope_id 的节点
    """
    from infrastructure.db.neo4j import get_driver

    try:
        # 查副本状态
        rows = await _query_scene(
            "SELECT * FROM ontol_cope_version WHERE id=? AND delete_flag='0'",
            (cope_id,),
        )
        if not rows:
            raise HTTPException(status_code=404, detail="Cope version not found")
        cope = rows[0]
        status = cope.get("cope_version_status", "00")

        driver = await get_driver()
        async with driver.session() as session:
            if status == "00":
                cypher = """
                    MATCH (n)
                    WHERE n.copy_version IS NULL OR n.copy_version = ''
                    OPTIONAL MATCH (n)-[r]-(m)
                    WHERE m.copy_version IS NULL OR m.copy_version = ''
                    RETURN n, collect(DISTINCT {edge: r, node: m}) AS rels
                    LIMIT 500
                """
            else:
                cypher = """
                    MATCH (n)
                    WHERE n.copy_version = $cope_id
                    OPTIONAL MATCH (n)-[r]-(m)
                    WHERE m.copy_version = $cope_id
                    RETURN n, collect(DISTINCT {edge: r, node: m}) AS rels
                    LIMIT 500
                """
            result = await session.run(cypher, cope_id=cope_id)
            records = [record async for record in result]

        nodes = {}
        edges = {}
        for record in records:
            n = record["n"]
            nid = n.element_id
            if nid not in nodes:
                nodes[nid] = {
                    "id": n.element_id,
                    "labels": list(n.labels),
                    "properties": dict(n.items()),
                }
            rels_list = record.get("rels")
            if rels_list:
                for rel_item in rels_list:
                    r = rel_item.get("edge")
                    m = rel_item.get("node")
                    if r is None or m is None:
                        continue
                    mid = m.element_id
                    if mid not in nodes:
                        nodes[mid] = {
                            "id": m.element_id,
                            "labels": list(m.labels),
                            "properties": dict(m.items()),
                        }
                    ekey = str(r.element_id)
                    if ekey not in edges:
                        edges[ekey] = {
                            "edge_id": r.element_id,
                            "source_id": r.start_node.element_id,
                            "target_id": r.end_node.element_id,
                            "type": r.type,
                            "properties": dict(r.items()),
                        }

        return {
            "cope_id": cope_id,
            "cope_version_status": status,
            "cope_name": cope.get("name", ""),
            "init_note_id": cope.get("init_note_id", ""),
            "init_note_name": cope.get("init_note_name", ""),
            "confidence": float(cope.get("confidence", 0.8)),
            "nodes": list(nodes.values()),
            "edges": list(edges.values()),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cope graph query failed: {e}")


# =========================================================================
# 对话-推演副本绑定 CRUD (ontol_chat_cope_version_relation)
# =========================================================================

class ChatCopeVersionBind(BaseModel):
    chat_id: str = Field(..., description="对话UUID")
    cope_version_id: str = Field(..., description="副本ID")


@router.post("/chat-cope-versions/bind")
async def bind_chat_cope_version(body: ChatCopeVersionBind):
    """将对话绑定到推演副本（先删旧绑定，再插入新绑定）。"""
    import uuid as _uuid
    rid = _uuid.uuid4().hex[:16]
    now = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    try:
        # 先清理旧绑定
        await _execute_scene(
            "UPDATE ontol_chat_cope_version_relation SET delete_flag='1' WHERE chat_id=?",
            (body.chat_id,),
        )
        # 插入新绑定
        await _execute_scene(
            "INSERT INTO ontol_chat_cope_version_relation (id, name, code, chat_id, cope_version_id, create_time) "
            "VALUES (?,?,?,?,?,?)",
            (rid, '', rid, body.chat_id, body.cope_version_id, now),
        )
        return {"id": rid, "chat_id": body.chat_id, "cope_version_id": body.cope_version_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Bind chat cope version failed: {e}")


@router.get("/chat-cope-versions/{chat_id}")
async def get_chat_cope_version(chat_id: str):
    """查询对话绑定的推演副本。"""
    try:
        rows = await _query_scene(
            "SELECT r.*, c.name AS cope_name, c.init_note_name, c.cope_version_status "
            "FROM ontol_chat_cope_version_relation r "
            "LEFT JOIN ontol_cope_version c ON r.cope_version_id = c.id "
            "WHERE r.chat_id=? AND r.delete_flag='0'",
            (chat_id,),
        )
        return rows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query chat cope version failed: {e}")


@router.delete("/chat-cope-versions/{relation_id}")
async def unbind_chat_cope_version(relation_id: str):
    """删除对话-副本绑定（软删除）。"""
    try:
        await _execute_scene(
            "UPDATE ontol_chat_cope_version_relation SET delete_flag='1' WHERE id=?",
            (relation_id,),
        )
        return {"deleted": True, "relation_id": relation_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unbind chat cope version failed: {e}")


# =========================================================================
# 审核记录 API — 薄路由，业务逻辑在 business/audit/audit_service.py
# =========================================================================


@router.get("/audit-logs")
async def list_audit_logs(
    audit_status: Optional[str] = None,
    trigger_source: Optional[str] = None,
    node_id: Optional[str] = None,
    batch_id: Optional[str] = None,
    keyword: Optional[str] = None,
    limit: int = 20,
    offset: int = 0,
):
    """分页查询审核记录。"""
    from business.audit.audit_service import list_audit_logs as _svc_list
    return _svc_list(
        audit_status=audit_status, trigger_source=trigger_source,
        node_id=node_id, batch_id=batch_id, keyword=keyword,
        limit=limit, offset=offset,
    )


@router.get("/audit-logs/{log_id}")
async def get_audit_log(log_id: str):
    """获取单条审核记录。"""
    from business.audit.audit_service import get_audit_log as _svc_get
    result = _svc_get(log_id)
    if result is None:
        raise HTTPException(status_code=404, detail="Audit log not found")
    return result


@router.post("/audit-logs")
async def create_audit_log(body: dict):
    """创建审核记录。"""
    from business.audit.audit_service import create_audit_log as _svc_create
    log_id = _svc_create(body)
    return {"id": log_id, "created": True}


@router.put("/audit-logs/{log_id}")
async def update_audit_log(log_id: str, body: dict):
    """更新审核记录（复核字段+状态）。"""
    from business.audit.audit_service import update_audit_log as _svc_update
    ok = _svc_update(log_id, body)
    if not ok:
        raise HTTPException(status_code=404, detail="Audit log not found")
    return {"updated": True, "id": log_id}


@router.delete("/audit-logs/{log_id}")
async def delete_audit_log(log_id: str):
    """软删除审核记录。"""
    from business.audit.audit_service import delete_audit_log as _svc_delete
    _svc_delete(log_id)
    return {"deleted": True, "id": log_id}
