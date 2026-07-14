"""文件解析 — docx/doc 文本提取 + LLM JSON 解析 + 两阶段解析管线。"""
import json
import os
import re
from pathlib import Path


def extract_text_from_docx(path: str) -> str:
    """从 .docx 文件提取纯文本（段落 + 表格）。"""
    from docx import Document

    doc = Document(path)
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    content = "\n\n".join(paragraphs)
    if not content.strip():
        raise ValueError("未从 .docx 文件中提取到文本内容")
    return content


def extract_text_from_doc(path: str) -> str:
    """从 .doc 文件提取纯文本（依赖 antiword 命令行工具）。"""
    import subprocess
    import shutil

    antiword = shutil.which("antiword")
    if not antiword:
        raise RuntimeError(
            "服务器未安装 antiword，无法解析 .doc 文件。"
            "请将文件另存为 .docx 或 .txt 格式后重新上传。"
        )

    result = subprocess.run(
        [antiword, path], capture_output=True, text=True, timeout=30,
    )
    if result.returncode != 0:
        stderr = result.stderr.strip()
        raise RuntimeError(f"antiword 解析 .doc 文件失败: {stderr or '未知错误'}")

    content = result.stdout.strip()
    if not content:
        raise ValueError("未从 .doc 文件中提取到文本内容")
    return content


def parse_entities_json(text: str) -> dict:
    """从 LLM 输出中提取 JSON 格式的实体和关系（多级降级策略）。"""
    text = text.strip()

    # 1. 直接 JSON 解析
    try:
        result = json.loads(text)
        if isinstance(result, dict):
            return result
    except (json.JSONDecodeError, TypeError):
        pass

    # 2. 提取 markdown 代码围栏
    fence_patterns = [
        r'```json\s*\n(.*?)\n\s*```',
        r'```\s*\n(.*?)\n\s*```',
        r'```json\s*(.*?)\s*```',
        r'```\s*(.*?)\s*```',
    ]
    for pattern in fence_patterns:
        for m in re.finditer(pattern, text, re.DOTALL | re.IGNORECASE):
            try:
                result = json.loads(m.group(1).strip())
                if isinstance(result, dict):
                    return result
            except (json.JSONDecodeError, TypeError):
                continue

    # 3. 按大括号平衡匹配
    brace_depth = 0
    start = -1
    candidates = []
    for i, ch in enumerate(text):
        if ch == '{':
            if brace_depth == 0:
                start = i
            brace_depth += 1
        elif ch == '}':
            brace_depth -= 1
            if brace_depth == 0 and start >= 0:
                candidates.append(text[start:i + 1])
                start = -1

    for candidate in candidates:
        try:
            result = json.loads(candidate)
            if isinstance(result, dict) and ('entities' in result or 'relationships' in result):
                return result
        except (json.JSONDecodeError, TypeError):
            continue

    # 4. 最终降级：去 Markdown 标记后贪婪匹配
    cleaned = re.sub(r'\*\*[^*]+\*\*', '', text)
    cleaned = re.sub(r'`[^`]+`', '', cleaned)
    m = re.search(r'\{.*\}', cleaned, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(0))
        except (json.JSONDecodeError, TypeError):
            pass

    return {
        "entities": [],
        "relationships": [],
        "_parse_error": True,
        "_raw_snippet": text[:500],
    }


async def _run_parse_pipeline(filename: str, model: str = "") -> dict:
    """两阶段 AI 解析管线：文本提取 → LLM 分类 → LLM 字段提取。

    Returns: {filename, entity_count, relationship_count, type_counts, entities, relationships, ...}
    """
    from business.upload.prompts import build_classify_prompt, build_extract_prompt
    from business.ontology import get_inherited_fields
    from capabilities.models.resolver import resolve_llm
    from langchain_core.messages import HumanMessage

    upload_dir = Path("infrastructure/storage/uploads")
    safe_name = os.path.basename(filename)
    file_path = upload_dir / safe_name

    if not file_path.exists():
        raise FileNotFoundError(f"File '{safe_name}' not found")

    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        content = extract_text_from_docx(str(file_path))
    elif suffix == ".doc":
        content = extract_text_from_doc(str(file_path))
    else:
        content = file_path.read_text(encoding="utf-8")

    llm_iface = resolve_llm(model)
    llm = await llm_iface.get_llm()

    CHUNK_SIZE = 3000
    chunks: list[str] = []
    start = 0
    while start < len(content):
        end = min(start + CHUNK_SIZE, len(content))
        if end < len(content):
            for sep in ("\n\n", "\n", "。", "；", ". "):
                pos = content.rfind(sep, start, end)
                if pos > start + 500:
                    end = pos + len(sep)
                    break
        chunks.append(content[start:end].strip())
        start = end

    chunk_errors: list[dict] = []

    # 阶段 1: 分类
    classify_prompt = build_classify_prompt()
    all_classified: dict[str, dict] = {}
    all_relationships: list[dict] = []
    seen_rels = set()

    for i, chunk in enumerate(chunks):
        if not chunk or len(chunk) < 10:
            continue
        try:
            response = await llm.ainvoke([
                HumanMessage(content=classify_prompt),
                HumanMessage(content=f"请解析以下文本，识别实体和关系（第{i+1}/{len(chunks)}块）：\n\n{chunk}"),
            ])
            text = response.content if hasattr(response, "content") else str(response)
            result = parse_entities_json(text)
            if result.pop("_parse_error", None):
                chunk_errors.append({"chunk_index": i+1, "reason": "分类阶段 JSON 解析失败"})
                continue
            for ent in result.get("entities", []):
                name = (ent.get("name") or "").strip()
                if name and name not in all_classified:
                    all_classified[name] = {"name": name, "ont_type": ent.get("ont_type", "M_ENTITY")}
            for rel in result.get("relationships", []):
                s = (rel.get("start_node_id") or rel.get("subject") or "").strip()
                p = (rel.get("type") or rel.get("predicate") or "").strip()
                o = (rel.get("end_node_id") or rel.get("object") or "").strip()
                if s and o and p:
                    key = f"{s}|{p}|{o}"
                    if key not in seen_rels:
                        seen_rels.add(key)
                        all_relationships.append({"start_node_id": s, "type": p, "end_node_id": o, "properties": rel.get("properties", {})})
        except Exception as e:
            chunk_errors.append({"chunk_index": i+1, "reason": f"分类阶段 LLM 调用失败: {str(e)}"})

    if not all_classified:
        return {"filename": safe_name, "entity_count": 0, "relationship_count": len(all_relationships),
                "type_counts": {}, "entities": [], "relationships": all_relationships,
                "phase1_classified": 0, "chunk_errors": chunk_errors[:10]}

    # 阶段 2: 按类型分组提取字段
    by_type: dict[str, list[str]] = {}
    for name, info in all_classified.items():
        by_type.setdefault(info["ont_type"], []).append(name)

    all_entities: dict[str, dict] = {}
    extract_errors: list[dict] = []

    for ont_type, names in by_type.items():
        inherited = get_inherited_fields(ont_type)
        if not inherited:
            for name in names:
                all_entities[name] = {"name": name, "ont_type": ont_type, "type_name": "", "properties": {}}
            continue

        extract_prompt = build_extract_prompt(ont_type)
        entity_list_text = "\n".join(f"- {n}" for n in names)
        try:
            response = await llm.ainvoke([
                HumanMessage(content=extract_prompt),
                HumanMessage(content=f"以下实体需要提取字段值（类型={ont_type}）：\n{entity_list_text}\n\n原始文本已在上文中提供，请为每个实体提取字段值，返回 JSON："),
            ])
            text = response.content if hasattr(response, "content") else str(response)
            result = parse_entities_json(text)
            if result.pop("_parse_error", None):
                extract_errors.append({"ont_type": ont_type, "reason": "字段提取 JSON 解析失败"})
                for name in names:
                    all_entities[name] = {"name": name, "ont_type": ont_type, "type_name": "", "properties": {}}
                continue
            for ent in result.get("entities", []):
                nm = (ent.get("name") or "").strip()
                if nm in names or nm in all_classified:
                    all_entities[nm] = {"name": nm, "ont_type": ont_type, "type_name": ent.get("type_name", ""), "properties": ent.get("properties", {})}
            for name in names:
                if name not in all_entities:
                    all_entities[name] = {"name": name, "ont_type": ont_type, "type_name": "", "properties": {}}
        except Exception as e:
            extract_errors.append({"ont_type": ont_type, "reason": f"字段提取 LLM 调用失败: {str(e)}"})
            for name in names:
                all_entities[name] = {"name": name, "ont_type": ont_type, "type_name": "", "properties": {}}

    entities_list = list(all_entities.values())
    type_counts = {}
    for e in entities_list:
        t = e.get("ont_type", "M_ENTITY")
        type_counts[t] = type_counts.get(t, 0) + 1

    valid_chunks = [c for c in chunks if c and len(c) >= 10]
    return {
        "filename": safe_name,
        "entity_count": len(entities_list),
        "relationship_count": len(all_relationships),
        "type_counts": type_counts,
        "entities": entities_list,
        "relationships": all_relationships,
        "chunks_total": len(valid_chunks),
        "chunks_ok": len(valid_chunks) - len(chunk_errors),
        "chunks_failed": len(chunk_errors),
        "chunk_errors": chunk_errors[:10],
        "extract_errors": extract_errors[:10],
    }
